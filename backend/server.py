from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import FileResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional
import uuid
from datetime import datetime, timezone, timedelta
from zoneinfo import ZoneInfo
from passlib.context import CryptContext
import jwt
import random
import shutil
import subprocess
from docx import Document
from emergentintegrations.llm.chat import LlmChat, UserMessage
import json
import asyncio

# Malaysian Timezone (UTC+8)
MALAYSIA_TZ = ZoneInfo("Asia/Kuala_Lumpur")

def get_malaysia_time():
    """Get current time in Malaysian timezone"""
    return datetime.now(MALAYSIA_TZ)

def get_malaysia_date():
    """Get current date in Malaysian timezone"""
    return get_malaysia_time().date()

def get_malaysia_time_str():
    """Get current time as string in HH:MM:SS format (Malaysian timezone)"""
    return get_malaysia_time().strftime("%H:%M:%S")

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db_name = os.environ.get('DB_NAME')
if not db_name:
    raise ValueError("DB_NAME environment variable is required")
db = client[db_name]
logging.info(f"🔥🔥🔥 CONNECTED TO DATABASE: {db_name} 🔥🔥🔥")

# Security
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
security = HTTPBearer()
SECRET_KEY = os.environ.get('SECRET_KEY')
if not SECRET_KEY or SECRET_KEY == 'your-secret-key-change-in-production':
    raise ValueError("SECRET_KEY environment variable must be set to a secure random value")
ALGORITHM = "HS256"

# Create the main app
app = FastAPI()
api_router = APIRouter(prefix="/api")

# Static files directory
STATIC_DIR = ROOT_DIR / "static"
STATIC_DIR.mkdir(exist_ok=True)
LOGO_DIR = STATIC_DIR / "logos"
LOGO_DIR.mkdir(exist_ok=True)
CERTIFICATE_DIR = STATIC_DIR / "certificates"
CERTIFICATE_DIR.mkdir(exist_ok=True)
CERTIFICATE_PDF_DIR = STATIC_DIR / "certificates_pdf"
CERTIFICATE_PDF_DIR.mkdir(exist_ok=True)
REPORT_DIR = STATIC_DIR / "reports"
REPORT_DIR.mkdir(exist_ok=True)
REPORT_PDF_DIR = STATIC_DIR / "reports_pdf"
REPORT_PDF_DIR.mkdir(exist_ok=True)
TEMPLATE_DIR = STATIC_DIR / "templates"
TEMPLATE_DIR.mkdir(exist_ok=True)
TEMPLATE_DIR = STATIC_DIR / "templates"
TEMPLATE_DIR.mkdir(exist_ok=True)
CHECKLIST_PHOTOS_DIR = STATIC_DIR / "checklist_photos"
CHECKLIST_PHOTOS_DIR.mkdir(exist_ok=True)

# ============ MODELS ============

class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: EmailStr
    full_name: str
    id_number: str
    role: str
    company_id: Optional[str] = None
    location: Optional[str] = None
    phone_number: Optional[str] = None
    created_at: datetime = Field(default_factory=get_malaysia_time)
    is_active: bool = True

class UserCreate(BaseModel):
    email: Optional[EmailStr] = None
    password: Optional[str] = None
    full_name: str
    id_number: str
    role: str
    company_id: Optional[str] = None
    location: Optional[str] = None
    phone_number: Optional[str] = None

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class TokenResponse(BaseModel):
    access_token: str
    token_type: str
    user: User

class Company(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    created_at: datetime = Field(default_factory=get_malaysia_time)

class CompanyCreate(BaseModel):
    name: str

class CompanyUpdate(BaseModel):
    name: str

class Program(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    description: Optional[str] = None
    pass_percentage: float = 70.0
    created_at: datetime = Field(default_factory=get_malaysia_time)

class ProgramCreate(BaseModel):
    name: str
    description: Optional[str] = None
    pass_percentage: Optional[float] = 70.0

class ProgramUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    pass_percentage: Optional[float] = None

class Session(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    program_id: str
    company_id: str
    location: str
    start_date: str
    end_date: str
    supervisor_ids: List[str] = []
    participant_ids: List[str] = []
    trainer_assignments: List[dict] = []
    coordinator_id: Optional[str] = None
    status: str = "active"  # "active" or "inactive"
    created_at: datetime = Field(default_factory=get_malaysia_time)

class ParticipantData(BaseModel):
    email: EmailStr
    password: str
    full_name: str
    id_number: str
    phone_number: Optional[str] = None

class SupervisorData(BaseModel):
    email: EmailStr
    password: str
    full_name: str
    id_number: str
    phone_number: Optional[str] = None

class SessionCreate(BaseModel):
    name: str
    program_id: str
    company_id: str
    location: str
    start_date: str
    end_date: str
    supervisor_ids: List[str] = []
    participant_ids: List[str] = []
    participants: List[ParticipantData] = []  # New participants to create or link
    supervisors: List[SupervisorData] = []  # New supervisors to create or link
    trainer_assignments: List[dict] = []
    coordinator_id: Optional[str] = None

class ParticipantAccess(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    participant_id: str
    session_id: str
    can_access_pre_test: bool = False
    can_access_post_test: bool = False
    can_access_checklist: bool = False
    can_access_feedback: bool = False
    pre_test_completed: bool = False
    post_test_completed: bool = False
    checklist_submitted: bool = False
    feedback_submitted: bool = False
    certificate_url: Optional[str] = None
    certificate_uploaded_at: Optional[str] = None
    certificate_uploaded_by: Optional[str] = None

class UpdateParticipantAccess(BaseModel):
    participant_id: str
    session_id: str
    can_access_pre_test: Optional[bool] = None
    can_access_post_test: Optional[bool] = None
    can_access_checklist: Optional[bool] = None
    can_access_feedback: Optional[bool] = None

class TestQuestion(BaseModel):
    question: str
    options: List[str]
    correct_answer: int

class Test(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    program_id: str
    test_type: str
    questions: List[TestQuestion] = []
    created_at: datetime = Field(default_factory=get_malaysia_time)

class TestCreate(BaseModel):
    program_id: str
    test_type: str
    questions: List[TestQuestion]

class TestResult(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    test_id: str
    participant_id: str
    session_id: str
    test_type: str
    answers: List[int] = []
    score: float = 0.0
    total_questions: int = 0
    correct_answers: int = 0
    passed: bool = False
    submitted_at: datetime = Field(default_factory=get_malaysia_time)
    question_indices: Optional[List[int]] = None  # Store original question order for shuffled tests

class TestSubmit(BaseModel):
    test_id: str
    session_id: str
    answers: List[int]
    question_indices: Optional[List[int]] = None  # Original question indices for shuffled tests

class ChecklistTemplate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    program_id: str
    items: List[str] = []
    created_at: datetime = Field(default_factory=get_malaysia_time)

class ChecklistTemplateCreate(BaseModel):
    program_id: str
    items: List[str]

class VehicleChecklist(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    participant_id: str
    session_id: str
    interval: str
    checklist_items: List[dict] = []
    submitted_at: datetime = Field(default_factory=get_malaysia_time)
    verified_by: Optional[str] = None
    verified_at: Optional[datetime] = None
    verification_status: str = "pending"

class ChecklistSubmit(BaseModel):
    session_id: str
    interval: str
    checklist_items: List[dict]

class ChecklistVerify(BaseModel):
    checklist_id: str
    status: str
    comments: Optional[str] = None

class VehicleDetails(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    participant_id: str
    session_id: str
    vehicle_model: str
    registration_number: str
    roadtax_expiry: str
    created_at: datetime = Field(default_factory=get_malaysia_time)

class VehicleDetailsSubmit(BaseModel):
    session_id: str
    vehicle_model: str
    registration_number: str
    roadtax_expiry: str

class TrainingReport(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    session_id: str
    coordinator_id: str
    group_photo: Optional[str] = None
    theory_photo_1: Optional[str] = None
    theory_photo_2: Optional[str] = None
    practical_photo_1: Optional[str] = None
    practical_photo_2: Optional[str] = None
    practical_photo_3: Optional[str] = None
    additional_notes: Optional[str] = None
    status: str = "draft"  # draft, submitted
    created_at: datetime = Field(default_factory=get_malaysia_time)
    submitted_at: Optional[datetime] = None

class TrainingReportCreate(BaseModel):
    session_id: str
    group_photo: Optional[str] = None
    theory_photo_1: Optional[str] = None
    theory_photo_2: Optional[str] = None
    practical_photo_1: Optional[str] = None
    practical_photo_2: Optional[str] = None
    practical_photo_3: Optional[str] = None
    additional_notes: Optional[str] = None
    status: str = "draft"

class Attendance(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    participant_id: str
    session_id: str
    date: str
    clock_in: Optional[str] = None
    clock_out: Optional[str] = None
    created_at: datetime = Field(default_factory=get_malaysia_time)

class AttendanceClockIn(BaseModel):
    session_id: str

class AttendanceClockOut(BaseModel):
    session_id: str

# Helper function to convert DOCX to PDF
def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """Convert DOCX to PDF using LibreOffice"""
    try:
        # Verify input file exists
        if not docx_path.exists():
            logging.error(f"DOCX file not found: {docx_path}")
            return False
        
        # Use LibreOffice in headless mode to convert DOCX to PDF
        result = subprocess.run([
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(pdf_path.parent),
            str(docx_path)
        ], check=True, capture_output=True, timeout=30)
        
        # Verify output file was created
        if not pdf_path.exists():
            logging.error(f"PDF file was not created: {pdf_path}")
            logging.error(f"LibreOffice output: {result.stdout.decode()}")
            logging.error(f"LibreOffice errors: {result.stderr.decode()}")
            return False
        
        return True
    except subprocess.TimeoutExpired:
        logging.error("PDF conversion timed out after 30 seconds")
        return False
    except subprocess.CalledProcessError as e:
        logging.error(f"LibreOffice conversion failed: {e.stderr.decode() if e.stderr else str(e)}")
        return False
    except Exception as e:
        logging.error(f"PDF conversion failed: {str(e)}")
        return False

class ChecklistItem(BaseModel):
    item: str
    status: str  # "good", "needs_repair"
    comments: str = ""
    photo_url: Optional[str] = None

class TrainerChecklistSubmit(BaseModel):
    participant_id: str
    session_id: str
    items: List[ChecklistItem]
    chief_trainer_comments: Optional[str] = None  # Only for chief trainers

class FeedbackQuestion(BaseModel):
    question: str
    type: str  # "rating" or "text"
    required: bool = True

class FeedbackTemplate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    program_id: str
    questions: List[FeedbackQuestion]
    created_at: datetime = Field(default_factory=get_malaysia_time)

class FeedbackTemplateCreate(BaseModel):
    program_id: str
    questions: List[FeedbackQuestion]

class CourseFeedback(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    participant_id: str
    session_id: str
    program_id: str
    responses: List[dict]  # [{"question": str, "answer": str/int}]
    submitted_at: datetime = Field(default_factory=get_malaysia_time)

class FeedbackSubmit(BaseModel):
    session_id: str
    program_id: str
    responses: List[dict]  # [{"question": str, "answer": str/int}]

class Certificate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    participant_id: str
    session_id: str
    program_name: str
    issue_date: datetime = Field(default_factory=get_malaysia_time)
    certificate_url: Optional[str] = None

class Settings(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = "app_settings"
    logo_url: Optional[str] = None
    company_name: str = "Malaysian Defensive Driving and Riding Centre Sdn Bhd"
    primary_color: str = "#3b82f6"
    secondary_color: str = "#6366f1"
    footer_text: str = ""
    certificate_template_url: Optional[str] = None
    max_certificate_file_size_mb: int = 5  # Max certificate file size in MB
    updated_at: datetime = Field(default_factory=get_malaysia_time)

class SettingsUpdate(BaseModel):
    company_name: Optional[str] = None
    primary_color: Optional[str] = None
    secondary_color: Optional[str] = None


# Coordinator and Chief Trainer Feedback Models
class CoordinatorFeedbackTemplate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = "coordinator_feedback_template"
    questions: List[dict] = [
        {
            "id": "training_smoothness",
            "question": "How smoothly did the training session run?",
            "type": "rating",
            "scale": 5
        },
        {
            "id": "participant_engagement",
            "question": "Rate the overall participant engagement level",
            "type": "rating",
            "scale": 5
        },
        {
            "id": "logistics",
            "question": "Were logistics (venue, equipment, timing) adequate?",
            "type": "rating",
            "scale": 5
        },
        {
            "id": "overall_observations",
            "question": "Please provide your overall observations about the training session",
            "type": "text"
        },
        {
            "id": "issues_identified",
            "question": "What issues or challenges were identified during the session?",
            "type": "text"
        },
        {
            "id": "recommendations",
            "question": "What are your recommendations for future sessions?",
            "type": "text"
        }
    ]
    updated_at: datetime = Field(default_factory=get_malaysia_time)

class ChiefTrainerFeedbackTemplate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = "chief_trainer_feedback_template"
    questions: List[dict] = [
        {
            "id": "pre_assessment",
            "question": "What were your observations from the pre-assessment?",
            "type": "text"
        },
        {
            "id": "theory_engagement",
            "question": "How engaged were participants during the theory session?",
            "type": "rating",
            "scale": 5
        },
        {
            "id": "practical_performance",
            "question": "Rate the overall practical session performance",
            "type": "rating",
            "scale": 5
        },
        {
            "id": "challenges",
            "question": "What challenges were encountered during training?",
            "type": "text"
        },
        {
            "id": "participant_dedication",
            "question": "Rate participant dedication and effort",
            "type": "rating",
            "scale": 5
        },
        {
            "id": "overall_impressions",
            "question": "Please share your overall impressions and recommendations",
            "type": "text"
        }
    ]
    updated_at: datetime = Field(default_factory=get_malaysia_time)

class CoordinatorFeedback(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    session_id: str
    coordinator_id: str
    responses: dict = {}
    submitted_at: datetime = Field(default_factory=get_malaysia_time)

class ChiefTrainerFeedback(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    session_id: str
    trainer_id: str
    responses: dict = {}
    submitted_at: datetime = Field(default_factory=get_malaysia_time)

class FeedbackTemplateUpdate(BaseModel):
    questions: List[dict]

    footer_text: Optional[str] = None
    max_certificate_file_size_mb: Optional[int] = None

# ============ HELPER FUNCTIONS ============

def hash_password(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)

def create_access_token(data: dict, expires_delta: timedelta = timedelta(days=7)):
    to_encode = data.copy()
    # JWT expiration should remain in UTC for standard compliance
    expire = datetime.now(timezone.utc) + expires_delta
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        token = credentials.credentials
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid token")
        
        user_doc = await db.users.find_one({"id": user_id}, {"_id": 0})
        if not user_doc:
            raise HTTPException(status_code=401, detail="User not found")
        
        if isinstance(user_doc.get('created_at'), str):
            user_doc['created_at'] = datetime.fromisoformat(user_doc['created_at'])
        
        return User(**user_doc)
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")

async def get_or_create_participant_access(participant_id: str, session_id: str):
    access_doc = await db.participant_access.find_one(
        {"participant_id": participant_id, "session_id": session_id},
        {"_id": 0}
    )
    
    if not access_doc:
        access_obj = ParticipantAccess(
            participant_id=participant_id,
            session_id=session_id
        )
        doc = access_obj.model_dump()
        await db.participant_access.insert_one(doc)
        return access_obj
    
    return ParticipantAccess(**access_doc)

async def find_or_create_user(user_data: dict, role: str, company_id: str) -> dict:
    """
    Find existing user by fullname OR email OR id_number (any match)
    If found: update the user with new data
    If not found: create new user
    Returns: user dict with 'is_existing' flag and user data
    """
    full_name = user_data.get("full_name")
    email = user_data.get("email")
    id_number = user_data.get("id_number")
    phone_number = user_data.get("phone_number")
    
    # Search for existing user by fullname OR email OR id_number
    query = {"$or": []}
    
    if full_name:
        query["$or"].append({"full_name": full_name})
    if email:
        query["$or"].append({"email": email})
    if id_number:
        query["$or"].append({"id_number": id_number})
    
    # If no fields provided, skip search
    if not query["$or"]:
        query = None
    
    existing_user = None
    if query:
        existing_user = await db.users.find_one(query, {"_id": 0})
    
    if existing_user:
        # User found - update with new data
        update_data = {
            "email": email,
            "id_number": user_data.get("id_number"),
            "phone_number": phone_number,
            "company_id": company_id,
        }
        # Remove None values
        update_data = {k: v for k, v in update_data.items() if v is not None}
        
        await db.users.update_one(
            {"id": existing_user["id"]},
            {"$set": update_data}
        )
        
        # Return updated user data
        updated_user = await db.users.find_one({"id": existing_user["id"]}, {"_id": 0})
        if isinstance(updated_user.get('created_at'), str):
            updated_user['created_at'] = datetime.fromisoformat(updated_user['created_at'])
        
        return {
            "is_existing": True,
            "user": User(**updated_user)
        }
    else:
        # User not found - create new
        # For participants: use default password 'mddrc1' if no password provided
        password = user_data.get("password")
        if role == "participant" and not password:
            password = "mddrc1"  # Default password for participants
        
        hashed_password = pwd_context.hash(password)
        
        # For participants: email defaults to IC@mddrc.com if not provided
        if role == "participant" and not email:
            email = f"{user_data.get('id_number')}@mddrc.com"
        
        new_user = User(
            email=email,
            full_name=full_name,
            id_number=user_data.get("id_number"),
            role=role,
            company_id=company_id,
            phone_number=phone_number
        )
        
        user_doc = new_user.model_dump()
        user_doc["created_at"] = user_doc["created_at"].isoformat()
        user_doc["password"] = hashed_password
        
        await db.users.insert_one(user_doc)
        
        return {
            "is_existing": False,
            "user": new_user
        }

# Training Report Models
class TrainingReport(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    session_id: str
    program_id: str
    company_id: str
    generated_by: str  # coordinator_id
    content: str  # Markdown content
    status: str  # "draft" or "published"
    created_at: datetime = Field(default_factory=get_malaysia_time)
    published_at: Optional[datetime] = None
    published_to_supervisors: List[str] = []  # List of supervisor IDs

class ReportGenerateRequest(BaseModel):
    session_id: str

class ReportUpdateRequest(BaseModel):
    content: str

# ============ ROUTES ============

@api_router.get("/")
async def root():
    return {"message": "Defensive Driving Training API"}

# Auth Routes
@api_router.post("/auth/register", response_model=User)
async def register_user(user_data: UserCreate, current_user: User = Depends(get_current_user)):
    # Role-based access control:
    # - Admins can create any user
    # - Coordinators can only create participants
    # - Assistant Admins can only create participants
    if current_user.role == "coordinator" or current_user.role == "assistant_admin":
        if user_data.role != "participant":
            raise HTTPException(status_code=403, detail="You can only create participants")
    elif current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Access denied")
    
    # For participants: use default credentials if not provided
    password = user_data.password
    email = user_data.email
    
    if user_data.role == "participant":
        # Default password: mddrc1
        if not password:
            password = "mddrc1"
        # Default email: IC@mddrc.com
        if not email:
            email = f"{user_data.id_number}@mddrc.com"
    
    # Check if user exists by email OR IC number
    existing = await db.users.find_one({
        "$or": [
            {"email": email},
            {"id_number": user_data.id_number}
        ]
    }, {"_id": 0})
    
    if existing:
        raise HTTPException(status_code=400, detail=f"User already exists with this {'email' if existing.get('email') == email else 'IC number'}")
    
    hashed_pw = hash_password(password)
    user_obj = User(
        email=email,
        full_name=user_data.full_name,
        id_number=user_data.id_number,
        role=user_data.role,
        company_id=user_data.company_id,
        location=user_data.location
    )
    
    doc = user_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    doc['password'] = hashed_pw
    
    await db.users.insert_one(doc)
    return user_obj

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(user_data: UserLogin):
    # Allow login with email OR IC number
    user_doc = await db.users.find_one({
        "$or": [
            {"email": user_data.email},
            {"id_number": user_data.email}  # Allow IC as username
        ]
    }, {"_id": 0})
    if not user_doc:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Check for both 'password' and 'hashed_password' field names
    password_hash = user_doc.get('password') or user_doc.get('hashed_password')
    if not password_hash:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if not verify_password(user_data.password, password_hash):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if not user_doc.get('is_active', True):
        raise HTTPException(status_code=401, detail="Account is inactive")
    
    token = create_access_token({"sub": user_doc['id']})
    
    if isinstance(user_doc.get('created_at'), str):
        user_doc['created_at'] = datetime.fromisoformat(user_doc['created_at'])
    
    user_doc.pop('password', None)
    user_doc.pop('hashed_password', None)
    user = User(**user_doc)
    
    return TokenResponse(access_token=token, token_type="bearer", user=user)

@api_router.get("/auth/me", response_model=User)
async def get_me(current_user: User = Depends(get_current_user)):
    return current_user

class ForgotPasswordRequest(BaseModel):
    email: EmailStr

class ResetPasswordRequest(BaseModel):
    email: EmailStr
    new_password: str

@api_router.post("/auth/forgot-password")
async def forgot_password(request: ForgotPasswordRequest):
    """
    Simple forgot password endpoint that checks if user exists
    In production, this would send an email with reset link
    """
    user_doc = await db.users.find_one({"email": request.email}, {"_id": 0})
    
    # Always return success to prevent email enumeration
    if not user_doc:
        return {"message": "If an account exists with this email, password reset instructions have been sent"}
    
    # For MVP: Return success message
    # In production: Generate token, send email with reset link
    return {"message": "If an account exists with this email, password reset instructions have been sent"}

@api_router.post("/auth/reset-password")
async def reset_password(request: ResetPasswordRequest):
    """
    Reset password for a user
    In production, this would require a valid reset token from email
    For MVP: Allow direct reset with email verification
    """
    user_doc = await db.users.find_one({"email": request.email}, {"_id": 0})
    
    if not user_doc:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Hash new password
    hashed_password = pwd_context.hash(request.new_password)
    
    # Update password
    await db.users.update_one(
        {"email": request.email},
        {"$set": {"password": hashed_password}}
    )
    
    return {"message": "Password reset successfully"}

# Company Routes
@api_router.post("/companies", response_model=Company)
async def create_company(company_data: CompanyCreate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can create companies")
    
    company_obj = Company(name=company_data.name)
    doc = company_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.companies.insert_one(doc)
    return company_obj

@api_router.get("/companies", response_model=List[Company])
async def get_companies(current_user: User = Depends(get_current_user)):
    companies = await db.companies.find({}, {"_id": 0}).to_list(1000)
    for company in companies:
        if isinstance(company.get('created_at'), str):
            company['created_at'] = datetime.fromisoformat(company['created_at'])
    return companies

@api_router.put("/companies/{company_id}", response_model=Company)
async def update_company(company_id: str, company_data: CompanyUpdate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can update companies")
    
    result = await db.companies.update_one(
        {"id": company_id},
        {"$set": company_data.model_dump()}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Company not found")
    
    company_doc = await db.companies.find_one({"id": company_id}, {"_id": 0})
    if isinstance(company_doc.get('created_at'), str):
        company_doc['created_at'] = datetime.fromisoformat(company_doc['created_at'])
    return Company(**company_doc)

@api_router.delete("/companies/{company_id}")
async def delete_company(company_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can delete companies")
    
    result = await db.companies.delete_one({"id": company_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Company not found")
    
    return {"message": "Company deleted successfully"}

# Program Routes
@api_router.post("/programs", response_model=Program)
async def create_program(program_data: ProgramCreate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can create programs")
    
    program_obj = Program(**program_data.model_dump())
    doc = program_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.programs.insert_one(doc)
    return program_obj

@api_router.get("/programs", response_model=List[Program])
async def get_programs(current_user: User = Depends(get_current_user)):
    programs = await db.programs.find({}, {"_id": 0}).to_list(1000)
    for program in programs:
        if isinstance(program.get('created_at'), str):
            program['created_at'] = datetime.fromisoformat(program['created_at'])
    return programs

@api_router.put("/programs/{program_id}", response_model=Program)
async def update_program(program_id: str, program_data: ProgramUpdate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can update programs")
    
    update_data = {k: v for k, v in program_data.model_dump().items() if v is not None}
    
    result = await db.programs.update_one(
        {"id": program_id},
        {"$set": update_data}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Program not found")
    
    program_doc = await db.programs.find_one({"id": program_id}, {"_id": 0})
    if isinstance(program_doc.get('created_at'), str):
        program_doc['created_at'] = datetime.fromisoformat(program_doc['created_at'])
    return Program(**program_doc)

@api_router.delete("/programs/{program_id}")
async def delete_program(program_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can delete programs")
    
    result = await db.programs.delete_one({"id": program_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Program not found")
    
    return {"message": "Program deleted successfully"}

# User Delete Route
@api_router.delete("/users/{user_id}")
async def delete_user(user_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can delete users")
    
    if user_id == current_user.id:
        raise HTTPException(status_code=400, detail="Cannot delete your own account")
    
    result = await db.users.delete_one({"id": user_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    
    return {"message": "User deleted successfully"}

# Check if user exists
@api_router.post("/users/check-exists")
async def check_user_exists(
    full_name: Optional[str] = None,
    email: Optional[str] = None,
    id_number: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    """Check if a user exists by fullname OR email OR id_number"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can check user existence")
    
    query = {"$or": []}
    
    if full_name:
        query["$or"].append({"full_name": full_name})
    if email:
        query["$or"].append({"email": email})
    if id_number:
        query["$or"].append({"id_number": id_number})
    
    if not query["$or"]:
        return {"exists": False, "user": None}
    
    existing_user = await db.users.find_one(query, {"_id": 0, "hashed_password": 0})
    
    if existing_user:
        if isinstance(existing_user.get('created_at'), str):
            existing_user['created_at'] = datetime.fromisoformat(existing_user['created_at'])
        return {
            "exists": True,
            "user": User(**existing_user)
        }
    
    return {"exists": False, "user": None}

# Session Routes
@api_router.post("/sessions")
async def create_session(session_data: SessionCreate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can create sessions")
    
    # Process new participants (find or create)
    processed_participant_ids = list(session_data.participant_ids)  # Start with existing IDs
    participant_results = []
    
    for participant_data in session_data.participants:
        result = await find_or_create_user(
            participant_data.model_dump(),
            role="participant",
            company_id=session_data.company_id
        )
        processed_participant_ids.append(result["user"].id)
        participant_results.append({
            "name": result["user"].full_name,
            "email": result["user"].email,
            "is_existing": result["is_existing"]
        })
    
    # Process new supervisors (find or create)
    processed_supervisor_ids = list(session_data.supervisor_ids)  # Start with existing IDs
    supervisor_results = []
    
    for supervisor_data in session_data.supervisors:
        result = await find_or_create_user(
            supervisor_data.model_dump(),
            role="pic_supervisor",
            company_id=session_data.company_id
        )
        processed_supervisor_ids.append(result["user"].id)
        supervisor_results.append({
            "name": result["user"].full_name,
            "email": result["user"].email,
            "is_existing": result["is_existing"]
        })
    
    # Create session with processed IDs
    session_obj = Session(
        name=session_data.name,
        program_id=session_data.program_id,
        company_id=session_data.company_id,
        location=session_data.location,
        start_date=session_data.start_date,
        end_date=session_data.end_date,
        participant_ids=processed_participant_ids,
        supervisor_ids=processed_supervisor_ids,
        trainer_assignments=session_data.trainer_assignments,
        coordinator_id=session_data.coordinator_id,
    )
    
    doc = session_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.sessions.insert_one(doc)
    
    # Create participant access records
    for participant_id in processed_participant_ids:
        await get_or_create_participant_access(participant_id, session_obj.id)
    
    return {
        "session": session_obj,
        "participant_results": participant_results,
        "supervisor_results": supervisor_results
    }

@api_router.get("/sessions", response_model=List[Session])
async def get_sessions(current_user: User = Depends(get_current_user)):
    # Non-admin users only see active sessions
    query = {}
    if current_user.role not in ["admin"]:
        query["status"] = "active"
    
    if current_user.role == "participant":
        query["participant_ids"] = current_user.id
        sessions = await db.sessions.find(query, {"_id": 0}).to_list(1000)
        
        # Auto-create participant_access records for each session
        for session in sessions:
            await get_or_create_participant_access(current_user.id, session['id'])
    elif current_user.role == "supervisor":
        query["supervisor_ids"] = current_user.id
        sessions = await db.sessions.find(query, {"_id": 0}).to_list(1000)
    else:
        sessions = await db.sessions.find(query, {"_id": 0}).to_list(1000)
    
    for session in sessions:
        if isinstance(session.get('created_at'), str):
            session['created_at'] = datetime.fromisoformat(session['created_at'])
    return sessions

@api_router.put("/sessions/{session_id}/toggle-status")
async def toggle_session_status(session_id: str, current_user: User = Depends(get_current_user)):
    """Toggle session between active and inactive (Admin only)"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can change session status")
    
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    new_status = "inactive" if session.get("status", "active") == "active" else "active"
    
    await db.sessions.update_one(
        {"id": session_id},
        {"$set": {"status": new_status}}
    )
    
    return {"message": f"Session marked as {new_status}", "status": new_status}

@api_router.get("/sessions/{session_id}", response_model=Session)
async def get_session(session_id: str, current_user: User = Depends(get_current_user)):
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    if isinstance(session.get('created_at'), str):
        session['created_at'] = datetime.fromisoformat(session['created_at'])
    
    return session

@api_router.get("/sessions/{session_id}/participants")
async def get_session_participants(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can view participants")
    
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    participants = []
    for participant_id in session['participant_ids']:
        user_doc = await db.users.find_one({"id": participant_id}, {"_id": 0, "password": 0})
        if user_doc:
            if isinstance(user_doc.get('created_at'), str):
                user_doc['created_at'] = datetime.fromisoformat(user_doc['created_at'])
            
            access = await get_or_create_participant_access(participant_id, session_id)
            
            participants.append({
                "user": user_doc,
                "access": access.model_dump()
            })
    
    return participants

@api_router.put("/sessions/{session_id}")
async def update_session(session_id: str, session_data: dict, current_user: User = Depends(get_current_user)):
    # Allow admins to update any session, coordinators can update sessions they're assigned to
    if current_user.role == "coordinator":
        # Check if coordinator is assigned to this session
        session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")
        if session.get("coordinator_id") != current_user.id:
            raise HTTPException(status_code=403, detail="You can only update sessions assigned to you")
    elif current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins and coordinators can update sessions")
    
    result = await db.sessions.update_one(
        {"id": session_id},
        {"$set": session_data}
    )
    
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Session not found")
    
    return {"message": "Session updated successfully"}

@api_router.delete("/sessions/{session_id}")
async def delete_session(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can delete sessions")
    
    result = await db.sessions.delete_one({"id": session_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Also delete related participant_access records
    await db.participant_access.delete_many({"session_id": session_id})
    
    return {"message": "Session deleted successfully"}

# Participant Access Routes
@api_router.post("/participant-access/update")
async def update_participant_access(access_data: UpdateParticipantAccess, current_user: User = Depends(get_current_user)):
    # Allow admins and coordinators to update access
    if current_user.role == "coordinator":
        # Verify coordinator is assigned to this session
        session = await db.sessions.find_one({"id": access_data.session_id}, {"_id": 0})
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")
        if session.get("coordinator_id") != current_user.id:
            raise HTTPException(status_code=403, detail="You can only manage access for sessions assigned to you")
    elif current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins and coordinators can update access")
    
    await get_or_create_participant_access(access_data.participant_id, access_data.session_id)
    
    update_fields = {}
    if access_data.can_access_pre_test is not None:
        update_fields['can_access_pre_test'] = access_data.can_access_pre_test
    if access_data.can_access_post_test is not None:
        update_fields['can_access_post_test'] = access_data.can_access_post_test
    if access_data.can_access_checklist is not None:
        update_fields['can_access_checklist'] = access_data.can_access_checklist
    if access_data.can_access_feedback is not None:
        update_fields['can_access_feedback'] = access_data.can_access_feedback
    
    await db.participant_access.update_one(
        {"participant_id": access_data.participant_id, "session_id": access_data.session_id},
        {"$set": update_fields}
    )
    
    return {"message": "Access updated successfully"}

@api_router.get("/participant-access/{session_id}")
async def get_my_access(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can check access")
    
    access = await get_or_create_participant_access(current_user.id, session_id)
    return access

@api_router.get("/participant-access/session/{session_id}")
async def get_session_access(session_id: str, current_user: User = Depends(get_current_user)):
    """Get all participant access records for a session (for coordinators/admins)"""
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Access denied")
    
    access_records = await db.participant_access.find({"session_id": session_id}, {"_id": 0}).to_list(1000)
    return access_records

@api_router.post("/participant-access/session/{session_id}/toggle")
async def toggle_session_access(session_id: str, access_data: dict, current_user: User = Depends(get_current_user)):
    """Toggle access for all participants in a session (coordinator/admin)"""
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Only coordinators and admins can control access")
    
    # Get session to find all participants
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    access_type = access_data.get("access_type")
    enabled = access_data.get("enabled", False)
    
    # Map access_type to field name
    field_mapping = {
        "pre_test": "can_access_pre_test",
        "post_test": "can_access_post_test",
        "feedback": "can_access_feedback",
        "checklist": "can_access_checklist"
    }
    
    if access_type not in field_mapping:
        raise HTTPException(status_code=400, detail="Invalid access type")
    
    field_name = field_mapping[access_type]
    
    # Update all participant access records for this session
    participant_ids = session.get("participant_ids", [])
    
    for participant_id in participant_ids:
        # Ensure access record exists
        await get_or_create_participant_access(participant_id, session_id)
        
        # Update the field
        await db.participant_access.update_one(
            {"participant_id": participant_id, "session_id": session_id},
            {"$set": {field_name: enabled}}
        )
    
    status_text = "enabled" if enabled else "disabled"
    return {"message": f"{access_type} access {status_text} for {len(participant_ids)} participants"}

# Coordinator Control Routes
@api_router.post("/sessions/{session_id}/release-pre-test")
async def release_pre_test(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Only admins and coordinators can release tests")
    
    # Get session to verify it exists
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Update all participant access records for this session
    result = await db.participant_access.update_many(
        {"session_id": session_id},
        {"$set": {"can_access_pre_test": True}}
    )
    
    return {"message": f"Pre-test released to {result.modified_count} participants"}

@api_router.post("/sessions/{session_id}/release-post-test")
async def release_post_test(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Only admins and coordinators can release tests")
    
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    result = await db.participant_access.update_many(
        {"session_id": session_id},
        {"$set": {"can_access_post_test": True}}
    )
    
    return {"message": f"Post-test released to {result.modified_count} participants"}

@api_router.post("/sessions/{session_id}/release-feedback")
async def release_feedback(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Only admins and coordinators can release feedback")
    
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    result = await db.participant_access.update_many(
        {"session_id": session_id},
        {"$set": {"can_access_feedback": True}}
    )
    
    return {"message": f"Feedback form released to {result.modified_count} participants"}

@api_router.get("/sessions/{session_id}/status")
async def get_session_status(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role not in ["admin", "coordinator", "trainer"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    # Get session
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Get all participant access records
    access_records = await db.participant_access.find({"session_id": session_id}, {"_id": 0}).to_list(1000)
    
    total_participants = len(access_records)
    pre_test_released = any(a.get('can_access_pre_test', False) for a in access_records)
    post_test_released = any(a.get('can_access_post_test', False) for a in access_records)
    feedback_released = any(a.get('can_access_feedback', False) for a in access_records)
    
    pre_test_completed = sum(1 for a in access_records if a.get('pre_test_completed', False))
    post_test_completed = sum(1 for a in access_records if a.get('post_test_completed', False))
    feedback_submitted = sum(1 for a in access_records if a.get('feedback_submitted', False))
    
    return {
        "session_id": session_id,
        "session_name": session.get('name', ''),
        "total_participants": total_participants,
        "pre_test": {
            "released": pre_test_released,
            "completed": pre_test_completed
        },
        "post_test": {
            "released": post_test_released,
            "completed": post_test_completed
        },
        "feedback": {
            "released": feedback_released,
            "submitted": feedback_submitted
        }
    }

@api_router.get("/sessions/{session_id}/results-summary")
async def get_results_summary(session_id: str, current_user: User = Depends(get_current_user)):
    # Check if user has permission (admin, coordinator, or chief trainer)
    if current_user.role not in ["admin", "coordinator"]:
        # Check if trainer is chief trainer for this session
        if current_user.role == "trainer":
            session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
            if not session:
                raise HTTPException(status_code=404, detail="Session not found")
            
            # Check if user is a chief trainer in this session
            is_chief = any(
                t.get('trainer_id') == current_user.id and t.get('role') == 'chief'
                for t in session.get('trainer_assignments', [])
            )
            if not is_chief:
                raise HTTPException(status_code=403, detail="Only chief trainers can view results")
        else:
            raise HTTPException(status_code=403, detail="Unauthorized")
    
    # Get all participants in the session
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    participant_ids = session.get('participant_ids', [])
    
    # Get participant details
    participants = await db.users.find(
        {"id": {"$in": participant_ids}},
        {"_id": 0, "password": 0}
    ).to_list(1000)
    
    # Get test results for all participants
    test_results = await db.test_results.find(
        {"session_id": session_id},
        {"_id": 0}
    ).to_list(1000)
    
    # Get feedback for all participants
    feedbacks = await db.course_feedback.find(
        {"session_id": session_id},
        {"_id": 0}
    ).to_list(1000)
    
    # Build summary
    summary = []
    for participant in participants:
        p_results = [r for r in test_results if r['participant_id'] == participant['id']]
        p_feedback = next((f for f in feedbacks if f['participant_id'] == participant['id']), None)
        
        pre_test = next((r for r in p_results if r['test_type'] == 'pre'), None)
        post_test = next((r for r in p_results if r['test_type'] == 'post'), None)
        
        summary.append({
            "participant": {
                "id": participant['id'],
                "name": participant['full_name'],
                "email": participant['email']
            },
            "pre_test": {
                "completed": pre_test is not None,
                "score": pre_test['score'] if pre_test else 0,
                "correct": pre_test['correct_answers'] if pre_test else 0,
                "total": pre_test['total_questions'] if pre_test else 0,
                "passed": pre_test['passed'] if pre_test else False,
                "result_id": pre_test['id'] if pre_test else None
            },
            "post_test": {
                "completed": post_test is not None,
                "score": post_test['score'] if post_test else 0,
                "correct": post_test['correct_answers'] if post_test else 0,
                "total": post_test['total_questions'] if post_test else 0,
                "passed": post_test['passed'] if post_test else False,
                "result_id": post_test['id'] if post_test else None
            },
            "feedback_submitted": p_feedback is not None
        })
    
    return {
        "session_id": session_id,
        "session_name": session.get('name', ''),
        "program_id": session.get('program_id', ''),
        "participants": summary
    }

# User Routes
@api_router.get("/users", response_model=List[User])
async def get_users(role: Optional[str] = None, current_user: User = Depends(get_current_user)):
    if current_user.role not in ["admin", "supervisor", "coordinator"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    query = {}
    if role:
        query["role"] = role
    
    users = await db.users.find(query, {"_id": 0, "password": 0}).to_list(1000)
    for user in users:
        if isinstance(user.get('created_at'), str):
            user['created_at'] = datetime.fromisoformat(user['created_at'])
    return users

@api_router.get("/users/{user_id}", response_model=User)
async def get_user(user_id: str, current_user: User = Depends(get_current_user)):
    # Allow access if: admin, supervisor, or the user themselves
    if current_user.role not in ["admin", "supervisor", "trainer", "coordinator"] and current_user.id != user_id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if isinstance(user.get('created_at'), str):
        user['created_at'] = datetime.fromisoformat(user['created_at'])
    
    return user

# Test Routes
@api_router.post("/tests", response_model=Test)
async def create_test(test_data: TestCreate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can create tests")
    
    test_obj = Test(**test_data.model_dump())
    doc = test_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.tests.insert_one(doc)
    return test_obj

@api_router.get("/tests/program/{program_id}", response_model=List[Test])
async def get_tests_by_program(program_id: str, current_user: User = Depends(get_current_user)):
    tests = await db.tests.find({"program_id": program_id}, {"_id": 0}).to_list(100)
    for test in tests:
        if isinstance(test.get('created_at'), str):
            test['created_at'] = datetime.fromisoformat(test['created_at'])
    return tests

@api_router.delete("/tests/{test_id}")
async def delete_test(test_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can delete tests")
    
    result = await db.tests.delete_one({"id": test_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Test not found")
    
    return {"message": "Test deleted successfully"}

@api_router.get("/sessions/{session_id}/tests/available")
async def get_available_tests(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can access this")
    
    # Get session
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Get participant access
    access = await get_or_create_participant_access(current_user.id, session_id)
    
    # Get tests for the session's program
    tests = await db.tests.find({"program_id": session['program_id']}, {"_id": 0}).to_list(10)
    
    available_tests = []
    for test in tests:
        if isinstance(test.get('created_at'), str):
            test['created_at'] = datetime.fromisoformat(test['created_at'])
        
        test_type = test['test_type']
        can_access = False
        is_completed = False
        
        if test_type == "pre":
            can_access = access.can_access_pre_test
            is_completed = access.pre_test_completed
        elif test_type == "post":
            can_access = access.can_access_post_test
            is_completed = access.post_test_completed
        
        if can_access and not is_completed:
            # Don't send correct answers to participant
            test_copy = test.copy()
            questions = test['questions'].copy()
            
            # Shuffle post-test questions
            if test_type == "post":
                random.shuffle(questions)
            
            test_copy['questions'] = [
                {
                    'question': q['question'],
                    'options': q['options']
                }
                for q in questions
            ]
            available_tests.append(test_copy)
    
    return available_tests

@api_router.get("/tests/{test_id}")
async def get_test(test_id: str, current_user: User = Depends(get_current_user)):
    test_doc = await db.tests.find_one({"id": test_id}, {"_id": 0})
    if not test_doc:
        raise HTTPException(status_code=404, detail="Test not found")
    
    if isinstance(test_doc.get('created_at'), str):
        test_doc['created_at'] = datetime.fromisoformat(test_doc['created_at'])
    
    # Make a copy of questions for shuffling
    questions = test_doc['questions'].copy()
    
    # Shuffle post-test questions for participants
    if current_user.role == "participant" and test_doc['test_type'] == "post":
        random.shuffle(questions)
    
    # Don't send correct answers to participants before submission
    if current_user.role == "participant":
        test_doc['questions'] = [
            {
                'question': q['question'],
                'options': q['options'],
                'original_index': test_doc['questions'].index(q)  # Track original position
            }
            for q in questions
        ]
    else:
        test_doc['questions'] = questions
    
    return test_doc

@api_router.post("/tests/submit", response_model=TestResult)
async def submit_test(submission: TestSubmit, current_user: User = Depends(get_current_user)):
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can submit tests")
    
    test_doc = await db.tests.find_one({"id": submission.test_id}, {"_id": 0})
    if not test_doc:
        raise HTTPException(status_code=404, detail="Test not found")
    
    program_doc = await db.programs.find_one({"id": test_doc['program_id']}, {"_id": 0})
    pass_percentage = program_doc.get('pass_percentage', 70.0) if program_doc else 70.0
    
    questions = test_doc['questions']
    
    # Ensure both are integers for comparison
    correct = 0
    for i, ans in enumerate(submission.answers):
        if i < len(questions):
            # If question_indices provided (shuffled test), use original index
            if submission.question_indices and i < len(submission.question_indices):
                original_idx = submission.question_indices[i]
            else:
                original_idx = i
            
            if original_idx < len(questions):
                submitted_answer = int(ans)
                correct_answer = int(questions[original_idx]['correct_answer'])
                if submitted_answer == correct_answer:
                    correct += 1
    
    score = (correct / len(questions)) * 100 if questions else 0
    passed = score >= pass_percentage
    
    result_obj = TestResult(
        test_id=submission.test_id,
        participant_id=current_user.id,
        session_id=submission.session_id,
        test_type=test_doc['test_type'],
        answers=submission.answers,
        score=score,
        total_questions=len(questions),
        correct_answers=correct,
        passed=passed,
        question_indices=submission.question_indices  # Store the shuffled order
    )
    
    doc = result_obj.model_dump()
    doc['submitted_at'] = doc['submitted_at'].isoformat()
    
    await db.test_results.insert_one(doc)
    
    update_field = 'pre_test_completed' if test_doc['test_type'] == 'pre' else 'post_test_completed'
    await db.participant_access.update_one(
        {"participant_id": current_user.id, "session_id": submission.session_id},
        {"$set": {update_field: True}}
    )
    
    return result_obj

@api_router.get("/tests/results/participant/{participant_id}", response_model=List[TestResult])
async def get_participant_results(participant_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role == "participant" and current_user.id != participant_id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    results = await db.test_results.find({"participant_id": participant_id}, {"_id": 0}).to_list(100)
    for result in results:
        if isinstance(result.get('submitted_at'), str):
            result['submitted_at'] = datetime.fromisoformat(result['submitted_at'])
    return results

@api_router.get("/tests/results/session/{session_id}")
async def get_session_test_results(session_id: str, current_user: User = Depends(get_current_user)):
    """Get all test results for a session (for coordinators/admins)"""
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Access denied")
    
    results = await db.test_results.find({"session_id": session_id}, {"_id": 0}).to_list(1000)
    
    for result in results:
        if isinstance(result.get('submitted_at'), str):
            result['submitted_at'] = datetime.fromisoformat(result['submitted_at'])
    
    return results

@api_router.get("/tests/results/{result_id}")
async def get_test_result_detail(result_id: str, current_user: User = Depends(get_current_user)):
    result = await db.test_results.find_one({"id": result_id}, {"_id": 0})
    if not result:
        raise HTTPException(status_code=404, detail="Test result not found")
    
    # Participants can only see their own results
    if current_user.role == "participant" and result['participant_id'] != current_user.id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    if isinstance(result.get('submitted_at'), str):
        result['submitted_at'] = datetime.fromisoformat(result['submitted_at'])
    
    # Get the test questions with correct answers
    test = await db.tests.find_one({"id": result['test_id']}, {"_id": 0})
    if test:
        questions = test['questions']
        
        # If question_indices exists (shuffled test), reorder questions to match participant's view
        if result.get('question_indices'):
            reordered_questions = []
            for idx in result['question_indices']:
                if idx < len(questions):
                    reordered_questions.append(questions[idx])
            result['test_questions'] = reordered_questions
        else:
            result['test_questions'] = questions
    
    return result

# Checklist Template Routes
@api_router.post("/checklist-templates", response_model=ChecklistTemplate)
async def create_checklist_template(template_data: ChecklistTemplateCreate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can create checklist templates")
    
    existing = await db.checklist_templates.find_one({"program_id": template_data.program_id}, {"_id": 0})
    if existing:
        await db.checklist_templates.update_one(
            {"program_id": template_data.program_id},
            {"$set": {"items": template_data.items}}
        )
        existing['items'] = template_data.items
        if isinstance(existing.get('created_at'), str):
            existing['created_at'] = datetime.fromisoformat(existing['created_at'])
        return ChecklistTemplate(**existing)
    
    template_obj = ChecklistTemplate(**template_data.model_dump())
    doc = template_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.checklist_templates.insert_one(doc)
    return template_obj

@api_router.get("/checklist-templates", response_model=List[ChecklistTemplate])
async def get_all_checklist_templates(current_user: User = Depends(get_current_user)):
    """Get all checklist templates"""
    templates = await db.checklist_templates.find({}, {"_id": 0}).to_list(length=None)
    result = []
    for template in templates:
        if isinstance(template.get('created_at'), str):
            template['created_at'] = datetime.fromisoformat(template['created_at'])
        result.append(ChecklistTemplate(**template))
    return result

@api_router.get("/checklist-templates/program/{program_id}", response_model=ChecklistTemplate)
async def get_checklist_template(program_id: str, current_user: User = Depends(get_current_user)):
    template = await db.checklist_templates.find_one({"program_id": program_id}, {"_id": 0})
    if not template:
        return ChecklistTemplate(program_id=program_id, items=[])
    
    if isinstance(template.get('created_at'), str):
        template['created_at'] = datetime.fromisoformat(template['created_at'])
    return ChecklistTemplate(**template)

@api_router.put("/checklist-templates/{template_id}", response_model=ChecklistTemplate)
async def update_checklist_template(template_id: str, template_data: ChecklistTemplateCreate, current_user: User = Depends(get_current_user)):
    """Update a checklist template"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can update checklist templates")
    
    existing = await db.checklist_templates.find_one({"id": template_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Template not found")
    
    await db.checklist_templates.update_one(
        {"id": template_id},
        {"$set": {"items": template_data.items, "program_id": template_data.program_id}}
    )
    
    existing['items'] = template_data.items
    existing['program_id'] = template_data.program_id
    if isinstance(existing.get('created_at'), str):
        existing['created_at'] = datetime.fromisoformat(existing['created_at'])
    
    return ChecklistTemplate(**existing)

@api_router.delete("/checklist-templates/{template_id}")
async def delete_checklist_template(template_id: str, current_user: User = Depends(get_current_user)):
    """Delete a checklist template"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can delete checklist templates")
    
    result = await db.checklist_templates.delete_one({"id": template_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Template not found")
    
    return {"message": "Template deleted successfully"}

# Vehicle Details Routes
@api_router.post("/vehicle-details/submit", response_model=VehicleDetails)
async def submit_vehicle_details(vehicle_data: VehicleDetailsSubmit, current_user: User = Depends(get_current_user)):
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can submit vehicle details")
    
    # Check if already exists
    existing = await db.vehicle_details.find_one({
        "participant_id": current_user.id,
        "session_id": vehicle_data.session_id
    }, {"_id": 0})
    
    if existing:
        # Update existing
        await db.vehicle_details.update_one(
            {"participant_id": current_user.id, "session_id": vehicle_data.session_id},
            {"$set": {
                "vehicle_model": vehicle_data.vehicle_model,
                "registration_number": vehicle_data.registration_number,
                "roadtax_expiry": vehicle_data.roadtax_expiry
            }}
        )
        existing.update(vehicle_data.model_dump())
        if isinstance(existing.get('created_at'), str):
            existing['created_at'] = datetime.fromisoformat(existing['created_at'])
        return VehicleDetails(**existing)
    
    vehicle_obj = VehicleDetails(
        participant_id=current_user.id,
        session_id=vehicle_data.session_id,
        vehicle_model=vehicle_data.vehicle_model,
        registration_number=vehicle_data.registration_number,
        roadtax_expiry=vehicle_data.roadtax_expiry
    )
    
    doc = vehicle_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.vehicle_details.insert_one(doc)
    return vehicle_obj

@api_router.get("/vehicle-details/{session_id}/{participant_id}")
async def get_vehicle_details(session_id: str, participant_id: str, current_user: User = Depends(get_current_user)):
    vehicle = await db.vehicle_details.find_one({
        "participant_id": participant_id,
        "session_id": session_id
    }, {"_id": 0})
    
    if not vehicle:
        return None
    
    if isinstance(vehicle.get('created_at'), str):
        vehicle['created_at'] = datetime.fromisoformat(vehicle['created_at'])
    return vehicle

# Attendance Routes
@api_router.post("/attendance/clock-in")
async def clock_in(attendance_data: AttendanceClockIn, current_user: User = Depends(get_current_user)):
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can clock in")
    
    # Use Malaysian time
    today = get_malaysia_date().isoformat()
    now = get_malaysia_time_str()
    
    # Check if already clocked in today
    existing = await db.attendance.find_one({
        "participant_id": current_user.id,
        "session_id": attendance_data.session_id,
        "date": today
    }, {"_id": 0})
    
    if existing and existing.get('clock_in'):
        raise HTTPException(status_code=400, detail="Already clocked in today")
    
    if existing:
        # Update existing
        await db.attendance.update_one(
            {"id": existing['id']},
            {"$set": {"clock_in": now}}
        )
        return {"message": "Clocked in successfully", "time": now}
    
    # Create new
    attendance_obj = Attendance(
        participant_id=current_user.id,
        session_id=attendance_data.session_id,
        date=today,
        clock_in=now
    )
    
    doc = attendance_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    await db.attendance.insert_one(doc)
    
    return {"message": "Clocked in successfully", "time": now}

@api_router.post("/attendance/clock-out")
async def clock_out(attendance_data: AttendanceClockOut, current_user: User = Depends(get_current_user)):
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can clock out")
    
    # Use Malaysian time
    today = get_malaysia_date().isoformat()
    now = get_malaysia_time_str()
    
    existing = await db.attendance.find_one({
        "participant_id": current_user.id,
        "session_id": attendance_data.session_id,
        "date": today
    }, {"_id": 0})
    
    if not existing or not existing.get('clock_in'):
        raise HTTPException(status_code=400, detail="Please clock in first")
    
    if existing.get('clock_out'):
        raise HTTPException(status_code=400, detail="Already clocked out today")
    
    await db.attendance.update_one(
        {"id": existing['id']},
        {"$set": {"clock_out": now}}
    )
    
    return {"message": "Clocked out successfully", "time": now}

@api_router.get("/attendance/session/{session_id}")
async def get_session_attendance(session_id: str, current_user: User = Depends(get_current_user)):
    """Get all attendance records for a session (for supervisors/coordinators)"""
    if current_user.role not in ["pic_supervisor", "coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Get session to verify access
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Get all attendance records for the session
    print(f"Querying attendance for session_id: {session_id}")
    logging.info(f"Querying attendance for session_id: {session_id}")
    attendance_records = await db.attendance.find({"session_id": session_id}, {"_id": 0}).to_list(1000)
    print(f"Found {len(attendance_records)} attendance records")
    logging.info(f"Found {len(attendance_records)} attendance records")
    
    # Get participant details only if we have attendance records
    participant_map = {}
    if attendance_records:
        participant_ids = list(set([r['participant_id'] for r in attendance_records]))
        logging.info(f"Looking up {len(participant_ids)} unique participants")
        
        if participant_ids:  # Only query if we have IDs to look up
            participants = await db.users.find({"id": {"$in": participant_ids}}, {"_id": 0}).to_list(1000)
            participant_map = {p['id']: p for p in participants}
            logging.info(f"Found {len(participants)} participant records")
    
    # Enrich attendance records with participant info
    for record in attendance_records:
        if isinstance(record.get('created_at'), str):
            record['created_at'] = datetime.fromisoformat(record['created_at'])
        participant = participant_map.get(record['participant_id'])
        if participant:
            record['participant_name'] = participant.get('full_name', 'Unknown')
            record['participant_email'] = participant.get('email', '')
        else:
            # Still include record even if participant not found
            record['participant_name'] = f"Participant {record['participant_id']}"
            record['participant_email'] = ''
            logging.warning(f"Could not find participant info for ID: {record['participant_id']}")
    
    return attendance_records

@api_router.get("/attendance/{session_id}/{participant_id}")
async def get_attendance(session_id: str, participant_id: str, current_user: User = Depends(get_current_user)):
    attendance_records = await db.attendance.find({
        "participant_id": participant_id,
        "session_id": session_id
    }, {"_id": 0}).to_list(100)
    
    for record in attendance_records:
        if isinstance(record.get('created_at'), str):
            record['created_at'] = datetime.fromisoformat(record['created_at'])
    
    return attendance_records
    
    print(f"=== ATTENDANCE ENDPOINT CALLED FOR SESSION: {session_id} ===")
    logging.info(f"=== ATTENDANCE ENDPOINT CALLED FOR SESSION: {session_id} ===")
    
    if current_user.role not in ["pic_supervisor", "coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Get session to verify access
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Get all attendance records for the session
    print(f"Querying attendance for session_id: {session_id}")
    logging.info(f"Querying attendance for session_id: {session_id}")
    attendance_records = await db.attendance.find({"session_id": session_id}, {"_id": 0}).to_list(1000)
    print(f"Found {len(attendance_records)} attendance records")
    logging.info(f"Found {len(attendance_records)} attendance records")
    
    # Get participant details only if we have attendance records
    participant_map = {}
    if attendance_records:
        participant_ids = list(set([r['participant_id'] for r in attendance_records]))
        logging.info(f"Looking up {len(participant_ids)} unique participants")
        
        if participant_ids:  # Only query if we have IDs to look up
            participants = await db.users.find({"id": {"$in": participant_ids}}, {"_id": 0}).to_list(1000)
            participant_map = {p['id']: p for p in participants}
            logging.info(f"Found {len(participants)} participant records")
    
    # Enrich attendance records with participant info
    for record in attendance_records:
        if isinstance(record.get('created_at'), str):
            record['created_at'] = datetime.fromisoformat(record['created_at'])
        participant = participant_map.get(record['participant_id'])
        if participant:
            record['participant_name'] = participant.get('full_name', 'Unknown')
            record['participant_email'] = participant.get('email', '')
        else:
            # Still include record even if participant not found
            record['participant_name'] = f"Participant {record['participant_id']}"
            record['participant_email'] = ''
            logging.warning(f"Could not find participant info for ID: {record['participant_id']}")
    
    return attendance_records

# Training Report Routes
@api_router.post("/training-reports", response_model=TrainingReport)
async def create_training_report(report_data: TrainingReportCreate, current_user: User = Depends(get_current_user)):
    """Create or update training completion report (coordinator only)"""
    if current_user.role != "coordinator":
        raise HTTPException(status_code=403, detail="Only coordinators can create training reports")
    
    # Check if report already exists for this session
    existing = await db.training_reports.find_one({"session_id": report_data.session_id}, {"_id": 0})
    
    if existing:
        # Update existing report
        update_data = report_data.model_dump()
        if update_data['status'] == 'submitted':
            update_data['submitted_at'] = get_malaysia_time().isoformat()
        
        await db.training_reports.update_one(
            {"session_id": report_data.session_id},
            {"$set": update_data}
        )
        
        updated = await db.training_reports.find_one({"session_id": report_data.session_id}, {"_id": 0})
        if isinstance(updated.get('created_at'), str):
            updated['created_at'] = datetime.fromisoformat(updated['created_at'])
        if isinstance(updated.get('submitted_at'), str):
            updated['submitted_at'] = datetime.fromisoformat(updated['submitted_at'])
        return TrainingReport(**updated)
    
    # Create new report
    report_obj = TrainingReport(
        **report_data.model_dump(),
        coordinator_id=current_user.id
    )
    
    if report_data.status == 'submitted':
        report_obj.submitted_at = datetime.now(timezone.utc)
    
    doc = report_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    if doc.get('submitted_at'):
        doc['submitted_at'] = doc['submitted_at'].isoformat()
    
    await db.training_reports.insert_one(doc)
    return report_obj

@api_router.get("/training-reports/{session_id}")
async def get_training_report(session_id: str, current_user: User = Depends(get_current_user)):
    """Get training report for a session"""
    report = await db.training_reports.find_one({"session_id": session_id}, {"_id": 0})
    
    if not report:
        # Return 404 instead of empty structure
        raise HTTPException(status_code=404, detail="Training report not found")
    
    if isinstance(report.get('created_at'), str):
        report['created_at'] = datetime.fromisoformat(report['created_at'])
    if isinstance(report.get('submitted_at'), str) and report.get('submitted_at'):
        report['submitted_at'] = datetime.fromisoformat(report['submitted_at'])
    
    return report

@api_router.get("/training-reports/coordinator/{coordinator_id}")
async def get_coordinator_reports(coordinator_id: str, current_user: User = Depends(get_current_user)):
    """Get all training reports for a coordinator"""
    if current_user.role != "coordinator" and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Access denied")
    
    reports = await db.training_reports.find({"coordinator_id": coordinator_id}, {"_id": 0}).to_list(100)
    
    for report in reports:
        if isinstance(report.get('created_at'), str):
            report['created_at'] = datetime.fromisoformat(report['created_at'])
        if isinstance(report.get('submitted_at'), str) and report.get('submitted_at'):
            report['submitted_at'] = datetime.fromisoformat(report['submitted_at'])
    
    return reports


@api_router.get("/training-reports/admin/all")
async def get_all_training_reports(
    search: Optional[str] = None,
    company_id: Optional[str] = None,
    program_id: Optional[str] = None,
    status: Optional[str] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    """Get all training reports with search and filter - Admin only"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access only")
    
    # Build query
    query = {"status": "submitted"}  # Only show submitted reports
    
    if status:
        query["status"] = status
    
    # Get all submitted reports
    reports = await db.training_reports.find(query, {"_id": 0}).to_list(1000)
    
    # Enrich each report with session, coordinator, company, program details
    enriched_reports = []
    
    for report in reports:
        session = await db.sessions.find_one({"id": report['session_id']}, {"_id": 0})
        if not session:
            continue
        
        # Get coordinator details
        coordinator = await db.users.find_one({"id": report.get('coordinator_id')}, {"_id": 0})
        
        # Get company and program details
        company = await db.companies.find_one({"id": session.get('company_id')}, {"_id": 0})
        program = await db.programs.find_one({"id": session.get('program_id')}, {"_id": 0})
        
        # Get participant count
        participant_count = len(session.get('participant_ids', []))
        
        # Apply filters
        if company_id and session.get('company_id') != company_id:
            continue
        
        if program_id and session.get('program_id') != program_id:
            continue
        
        # Apply date filter
        if start_date:
            session_date = session.get('start_date')
            if session_date and session_date < start_date:
                continue
        
        if end_date:
            session_date = session.get('end_date')
            if session_date and session_date > end_date:
                continue
        
        # Build enriched report
        enriched = {
            **report,
            "session_name": session.get('name', 'Unknown'),
            "session_start_date": session.get('start_date'),
            "session_end_date": session.get('end_date'),
            "session_location": session.get('location'),
            "coordinator_name": coordinator.get('full_name') if coordinator else 'Unknown',
            "company_name": company.get('name') if company else 'Unknown',
            "company_id": session.get('company_id'),
            "program_name": program.get('name') if program else 'Unknown',
            "program_id": session.get('program_id'),
            "participant_count": participant_count
        }
        
        # Apply search filter
        if search:
            search_lower = search.lower()
            searchable_text = f"{enriched['session_name']} {enriched['coordinator_name']} {enriched['company_name']} {enriched['program_name']} {enriched['session_location']}".lower()
            
            if search_lower not in searchable_text:
                continue
        
        enriched_reports.append(enriched)
    
    # Sort by submitted date (most recent first)
    enriched_reports.sort(key=lambda x: x.get('submitted_at', ''), reverse=True)
    
    return {
        "total": len(enriched_reports),
        "reports": enriched_reports
    }


@api_router.post("/training-reports/{session_id}/generate-ai-report")
async def generate_ai_report(session_id: str, current_user: User = Depends(get_current_user)):
    """Generate AI training report using ChatGPT"""
    if current_user.role != "coordinator" and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only coordinators can generate reports")
    
    from emergentintegrations.llm.chat import LlmChat, UserMessage, FileContentWithMimeType
    from dotenv import load_dotenv
    load_dotenv()
    
    # Get session details
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Get program details
    program = await db.programs.find_one({"id": session['program_id']}, {"_id": 0})
    
    # Get company details
    company = await db.companies.find_one({"id": session['company_id']}, {"_id": 0})
    
    # Get participants count
    participant_count = len(session.get('participant_ids', []))
    
    # Get attendance records
    attendance_records = await db.attendance.find({"session_id": session_id}, {"_id": 0}).to_list(1000)
    total_attendance = len(set([r['participant_id'] for r in attendance_records]))
    
    # Get test results
    test_results = await db.test_results.find({"session_id": session_id}, {"_id": 0}).to_list(1000)
    passed_tests = len([r for r in test_results if r.get('passed', False)])
    
    # Get training report with photos
    training_report = await db.training_reports.find_one({"session_id": session_id}, {"_id": 0})
    
    # Build context for AI
    context = f"""
Generate a professional defensive driving training completion report in a structured format similar to official training documentation.

**SESSION INFORMATION:**
Program Name: {program.get('name', 'N/A') if program else 'N/A'}
Company: {company.get('name', 'N/A') if company else 'N/A'}
Training Location: {session.get('location', 'N/A')}
Training Period: {session.get('start_date', 'N/A')} to {session.get('end_date', 'N/A')}
Total Participants: {participant_count}
Attendance: {total_attendance} out of {participant_count} participants
Assessment Pass Rate: {passed_tests} out of {len(test_results)} passed

**DOCUMENTATION:**
- Group Photo: {'Attached' if training_report and training_report.get('group_photo') else 'Not provided'}
- Theory Session Photos: {2 if training_report and training_report.get('theory_photo_1') and training_report.get('theory_photo_2') else 0} photos attached
- Practical Session Photos: {3 if training_report and training_report.get('practical_photo_1') and training_report.get('practical_photo_2') and training_report.get('practical_photo_3') else 0} photos attached

**REQUIRED REPORT STRUCTURE:**

# TRAINING COMPLETION REPORT

## 1. EXECUTIVE SUMMARY
[Provide a 2-3 sentence overview of the training session]

## 2. TRAINING PROGRAM DETAILS
- Program Name: [name]
- Training Duration: [dates]
- Location: [location]
- Target Audience: [company employees]

## 3. TRAINING OBJECTIVES
[List 3-4 key objectives of the defensive driving program]

## 4. TRAINING DELIVERY
**Theory Sessions:**
[Describe theory topics covered - 2-3 sentences]

**Practical Sessions:**
[Describe hands-on activities and exercises - 2-3 sentences]

## 5. PARTICIPANT PERFORMANCE
- Total Enrolled: {participant_count}
- Attendance Rate: {round((total_attendance/participant_count)*100) if participant_count > 0 else 0}%
- Assessment Pass Rate: {round((passed_tests/len(test_results))*100) if len(test_results) > 0 else 0}%

## 6. KEY LEARNING OUTCOMES
[List 4-5 key skills/knowledge participants gained]
- 
- 
- 

## 7. TRAINING EFFECTIVENESS
[Evaluate based on attendance and pass rates - 2-3 sentences]

## 8. OBSERVATIONS & FEEDBACK
[Note any significant observations about participant engagement, questions asked, areas of difficulty]

## 9. RECOMMENDATIONS
[Provide 2-3 recommendations for future training sessions]

## 10. CONCLUSION
[Summarize the overall success of the training]

---
Report Prepared By: Training Coordinator
Date: {get_malaysia_time().strftime('%Y-%m-%d')}

Please generate this report professionally with proper formatting, specific details based on the data provided, and maintain a formal tone suitable for official documentation.
"""
    
    try:
        # Initialize LLM Chat
        api_key = os.environ.get('EMERGENT_LLM_KEY', '')
        if not api_key:
            raise HTTPException(status_code=500, detail="EMERGENT_LLM_KEY not configured")
        
        chat = LlmChat(
            api_key=api_key,
            session_id=f"report_{session_id}",
            system_message="You are a professional training report writer specializing in defensive driving and road safety training programs."
        ).with_model("openai", "gpt-4o")
        
        user_message = UserMessage(text=context)
        
        # Generate report
        ai_response = await chat.send_message(user_message)
        
        return {
            "session_id": session_id,
            "generated_report": ai_response,
            "metadata": {
                "participant_count": participant_count,
                "attendance_rate": f"{total_attendance}/{participant_count}",
                "test_pass_rate": f"{passed_tests}/{len(test_results)}",
                "photos_included": bool(training_report)
            }
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate AI report: {str(e)}")


# Professional DOCX Report Generation
@api_router.post("/training-reports/{session_id}/generate-docx")
async def generate_docx_report(session_id: str, current_user: User = Depends(get_current_user)):
    """Generate a professional DOCX training report with all data populated"""
    
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Only coordinators and admins can generate reports")
    
    try:
        # Gather all session data
        session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")
        
        program = await db.programs.find_one({"id": session.get('program_id')}, {"_id": 0}) if session.get('program_id') else None
        company = await db.companies.find_one({"id": session.get('company_id')}, {"_id": 0}) if session.get('company_id') else None
        
        # Validate required data
        if not program:
            raise HTTPException(status_code=400, detail="Program not found for this session. Please ensure the session has a valid program assigned.")
        if not company:
            raise HTTPException(status_code=400, detail="Company not found for this session. Please ensure the session has a valid company assigned.")
        
        # Get participants with full details
        participant_ids = session.get('participant_ids', [])
        participants = []
        for pid in participant_ids:
            user = await db.users.find_one({"id": pid}, {"_id": 0})
            if user:
                # Get pre and post test results
                pre_test = await db.test_results.find_one({
                    "participant_id": pid,
                    "session_id": session_id,
                    "test_type": "pre"
                }, {"_id": 0})
                
                post_test = await db.test_results.find_one({
                    "participant_id": pid,
                    "session_id": session_id,
                    "test_type": "post"
                }, {"_id": 0})
                
                participants.append({
                    "name": user.get('full_name'),
                    "id_number": user.get('id_number', 'N/A'),
                    "pre_test_score": pre_test.get('score', 0) if pre_test else 0,
                    "pre_test_passed": pre_test.get('passed', False) if pre_test else False,
                    "post_test_score": post_test.get('score', 0) if post_test else 0,
                    "post_test_passed": post_test.get('passed', False) if post_test else False,
                    "improvement": (post_test.get('score', 0) if post_test else 0) - (pre_test.get('score', 0) if pre_test else 0)
                })
        
        # Get vehicle checklists with issues
        checklists = await db.vehicle_checklists.find({"session_id": session_id}, {"_id": 0}).to_list(100)
        vehicle_issues = []
        for checklist in checklists:
            participant = await db.users.find_one({"id": checklist['participant_id']}, {"_id": 0})
            issues_list = []
            for item in checklist.get('checklist_items', []):
                if item.get('status') == 'needs_repair':
                    issues_list.append({
                        "item": item.get('item', 'Unknown'),
                        "comment": item.get('comments', 'No comment'),
                        "photo_url": item.get('photo_url', '')
                    })
            
            if issues_list:
                vehicle_issues.append({
                    "participant_name": participant.get('full_name') if participant else 'Unknown',
                    "issues": issues_list
                })
        
        # Get training photos from training report
        training_report = await db.training_reports.find_one({"session_id": session_id}, {"_id": 0})
        training_photos = {
            "group_photo": training_report.get('group_photo') if training_report else None,
            "theory_photo_1": training_report.get('theory_photo_1') if training_report else None,
            "theory_photo_2": training_report.get('theory_photo_2') if training_report else None,
            "practical_photo_1": training_report.get('practical_photo_1') if training_report else None,
            "practical_photo_2": training_report.get('practical_photo_2') if training_report else None,
            "practical_photo_3": training_report.get('practical_photo_3') if training_report else None
        }
        
        # Get participant feedback
        all_feedback = await db.course_feedback.find({"session_id": session_id}, {"_id": 0}).to_list(100)
        feedback_data = []
        for feedback in all_feedback:
            participant = await db.users.find_one({"id": feedback['participant_id']}, {"_id": 0})
            feedback_data.append({
                "participant_name": participant.get('full_name') if participant else 'Unknown',
                "responses": feedback.get('responses', [])
            })
        
        # Determine vehicle type from program name for objectives
        program_name_lower = program.get('name', '').lower()
        is_motorcycle = 'motor' in program_name_lower or 'bike' in program_name_lower or 'rider' in program_name_lower
        is_truck = 'truck' in program_name_lower or 'lorry' in program_name_lower or 'heavy' in program_name_lower
        
        # Create DOCX document with enhanced formatting
        doc = Document()
        
        # COVER PAGE
        title = doc.add_heading('DEFENSIVE DRIVING/RIDING TRAINING', 0)
        title.alignment = 1  # Center alignment
        subtitle = doc.add_heading('COMPREHENSIVE COMPLETION REPORT', 0)
        subtitle.alignment = 1
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Cover details in a cleaner format
        cover_table = doc.add_table(rows=7, cols=2)
        cover_table.style = 'Light List Accent 1'
        cover_details = [
            ('Program:', program.get('name', 'N/A')),
            ('Company:', company.get('name', 'N/A')),
            ('Location:', session.get('location', 'N/A')),
            ('Training Period:', f"{session.get('start_date', 'N/A')} to {session.get('end_date', 'N/A')}"),
            ('Participants:', str(len(participants))),
            ('Submitted by:', current_user.full_name),
            ('Date:', get_malaysia_time().strftime('%Y-%m-%d'))
        ]
        for idx, (label, value) in enumerate(cover_details):
            cover_table.rows[idx].cells[0].text = label
            cover_table.rows[idx].cells[1].text = value
        
        doc.add_paragraph()
        doc.add_paragraph()
        footer_text = doc.add_paragraph('Prepared by: MDDRC (Malaysian Defensive Driving & Riding Centre)')
        footer_text.alignment = 1
        doc.add_page_break()
        
        # EXECUTIVE SUMMARY - COMPREHENSIVE
        doc.add_heading('1. EXECUTIVE SUMMARY', 1)
        pre_avg = sum([p['pre_test_score'] for p in participants]) / len(participants) if participants else 0
        post_avg = sum([p['post_test_score'] for p in participants]) / len(participants) if participants else 0
        improvement = post_avg - pre_avg
        
        # Count pass/fail statistics
        pre_pass_count = sum([1 for p in participants if p['pre_test_passed']])
        post_pass_count = sum([1 for p in participants if p['post_test_passed']])
        improved_count = sum([1 for p in participants if p['improvement'] > 0])
        
        doc.add_paragraph(
            f"This comprehensive report documents the Defensive {'Riding' if is_motorcycle else 'Driving'} Training "
            f"conducted for {company.get('name', 'N/A')} at {session.get('location', 'N/A')} from "
            f"{session.get('start_date', 'N/A')} to {session.get('end_date', 'N/A')}. The program was designed to "
            f"enhance safety awareness, reinforce defensive {'riding' if is_motorcycle else 'driving'} techniques, "
            f"and reduce commuting-related accidents, aligning with the company's commitment to employee safety."
        )
        doc.add_paragraph()
        
        doc.add_paragraph(
            f"The training program successfully engaged {len(participants)} participants through a structured "
            f"curriculum combining theoretical instruction and practical hands-on sessions. Participants demonstrated "
            f"high engagement levels and openness to feedback, contributing to a positive learning environment."
        )
        doc.add_paragraph()
        
        # KEY OUTCOMES heading
        doc.add_paragraph("KEY OUTCOMES:", style='Heading 3')
        outcomes = [
            f"• Total Participants: {len(participants)}",
            f"• Pre-Training Assessment Average: {pre_avg:.1f}%",
            f"• Post-Training Assessment Average: {post_avg:.1f}%",
            f"• Overall Improvement: {improvement:+.1f}%",
            f"• Pre-Test Pass Rate: {pre_pass_count}/{len(participants)} ({(pre_pass_count/len(participants)*100):.0f}% if len(participants) > 0 else 0)",
            f"• Post-Test Pass Rate: {post_pass_count}/{len(participants)} ({(post_pass_count/len(participants)*100):.0f}% if len(participants) > 0 else 0)",
            f"• Participants Showing Improvement: {improved_count}/{len(participants)} ({(improved_count/len(participants)*100):.0f}% if len(participants) > 0 else 0)"
        ]
        for outcome in outcomes:
            doc.add_paragraph(outcome)
        doc.add_paragraph()
        
        # TRAINING IMPACT
        doc.add_paragraph("TRAINING IMPACT:", style='Heading 3')
        doc.add_paragraph(
            f"The training successfully enhanced participants' understanding of hazard awareness, proper braking control, "
            f"and {'balance techniques' if is_motorcycle else 'vehicle control'}. Participants demonstrated improved ability "
            f"to identify potential road hazards and apply defensive {'riding' if is_motorcycle else 'driving'} principles. "
            f"The program fostered a culture of safety discipline and mutual learning among participants."
        )
        doc.add_paragraph()
        
        # SAFETY OBSERVATIONS (if vehicle issues found)
        if vehicle_issues:
            doc.add_paragraph("SAFETY OBSERVATIONS:", style='Heading 3')
            doc.add_paragraph(
                f"Vehicle inspections revealed {len(vehicle_issues)} {'motorcycles' if is_motorcycle else 'vehicles'} "
                f"with safety concerns requiring immediate attention. Detailed recommendations for addressing these issues "
                f"are provided in Section 9 of this report."
            )
        
        doc.add_page_break()
        
        # TRAINING OBJECTIVES
        doc.add_heading('2. TRAINING OBJECTIVES', 1)
        doc.add_paragraph(
            "This training program was designed with the following core objectives to enhance workplace safety and reduce accident risks:"
        )
        doc.add_paragraph()
        
        if is_motorcycle:
            objectives = [
                "• Improve rider safety awareness and hazard recognition on Malaysian roads",
                "• Reinforce defensive riding techniques for daily commuting",
                "• Reduce motorcycle-related accidents and injuries among employees",
                "• Promote proper Personal Protective Equipment (PPE) usage and motorcycle maintenance",
                "• Align riding behavior with company safety values and policies",
                "• Develop emergency response skills for critical road situations"
            ]
        elif is_truck:
            objectives = [
                "• Enhance heavy vehicle safety awareness and load management",
                "• Reinforce defensive driving techniques for commercial vehicles",
                "• Reduce delivery delays caused by accidents and vehicle breakdowns",
                "• Improve vehicle pre-trip inspection and maintenance practices",
                "• Minimize company liability and insurance costs through safer driving",
                "• Align driving behavior with company safety standards and regulations"
            ]
        else:  # Car/general driving
            objectives = [
                "• Improve driver safety awareness and hazard perception",
                "• Reinforce defensive driving techniques for daily operations",
                "• Reduce vehicle-related accidents and associated costs",
                "• Promote proper vehicle maintenance and pre-drive safety checks",
                "• Align driving behavior with company safety policies",
                "• Develop emergency response and accident avoidance skills"
            ]
        
        for objective in objectives:
            doc.add_paragraph(objective)
        doc.add_paragraph()
        doc.add_paragraph(
            "These objectives support the organization's commitment to employee welfare and operational excellence "
            "through enhanced road safety practices."
        )
        doc.add_page_break()
        
        # TRAINING AGENDA
        doc.add_heading('3. TRAINING AGENDA', 1)
        doc.add_paragraph(
            f"The training was conducted over a {2 if is_motorcycle else 2}-day period, combining theoretical instruction "
            f"with practical hands-on sessions:"
        )
        doc.add_paragraph()
        
        # DAY 1
        doc.add_heading('DAY 1 - Theory & Foundation', 2)
        if is_motorcycle:
            day1_items = [
                ('08:00 - 08:30', 'Registration & Welcome Briefing'),
                ('08:30 - 10:00', 'Hazard Recognition & Road Awareness'),
                ('10:00 - 10:15', 'Break'),
                ('10:15 - 12:00', 'Safe Distance Management & Speed Control'),
                ('12:00 - 13:00', 'Lunch'),
                ('13:00 - 14:30', 'Traffic Law & Regulations Review'),
                ('14:30 - 14:45', 'Break'),
                ('14:45 - 16:30', 'Fatigue Management & Weather Conditions'),
                ('16:30 - 17:00', 'Pre-Test Assessment & Day 1 Review')
            ]
        else:
            day1_items = [
                ('08:00 - 08:30', 'Registration & Welcome Briefing'),
                ('08:30 - 10:00', 'Defensive Driving Principles & Hazard Recognition'),
                ('10:00 - 10:15', 'Break'),
                ('10:15 - 12:00', 'Safe Following Distance & Speed Management'),
                ('12:00 - 13:00', 'Lunch'),
                ('13:00 - 14:30', 'Traffic Law & Road Safety Regulations'),
                ('14:30 - 14:45', 'Break'),
                ('14:45 - 16:30', 'Driver Fatigue & Weather Driving Conditions'),
                ('16:30 - 17:00', 'Pre-Test Assessment & Day 1 Summary')
            ]
        
        agenda_table_day1 = doc.add_table(rows=len(day1_items)+1, cols=2)
        agenda_table_day1.style = 'Light Grid Accent 1'
        agenda_table_day1.rows[0].cells[0].text = 'Time'
        agenda_table_day1.rows[0].cells[1].text = 'Activity'
        for idx, (time, activity) in enumerate(day1_items, 1):
            agenda_table_day1.rows[idx].cells[0].text = time
            agenda_table_day1.rows[idx].cells[1].text = activity
        
        doc.add_paragraph()
        
        # DAY 2
        doc.add_heading('DAY 2 - Practical Skills & Assessment', 2)
        if is_motorcycle:
            day2_items = [
                ('08:00 - 08:30', 'Day 2 Safety Briefing & PPE Check'),
                ('08:30 - 10:00', 'Emergency Braking Techniques (Practical)'),
                ('10:00 - 10:15', 'Break'),
                ('10:15 - 12:00', 'Obstacle Avoidance & Swerving Maneuvers'),
                ('12:00 - 13:00', 'Lunch'),
                ('13:00 - 14:30', 'Cornering Techniques & Body Positioning'),
                ('14:30 - 14:45', 'Break'),
                ('14:45 - 16:00', 'Left Lane Riding & Traffic Integration'),
                ('16:00 - 16:45', 'Post-Test Assessment'),
                ('16:45 - 17:00', 'Certificate Presentation & Closing')
            ]
        else:
            day2_items = [
                ('08:00 - 08:30', 'Day 2 Safety Briefing & Vehicle Check'),
                ('08:30 - 10:00', 'Emergency Braking & Stopping Techniques'),
                ('10:00 - 10:15', 'Break'),
                ('10:15 - 12:00', 'Obstacle Avoidance & Lane Change Maneuvers'),
                ('12:00 - 13:00', 'Lunch'),
                ('13:00 - 14:30', 'Cornering & Vehicle Control Exercises'),
                ('14:30 - 14:45', 'Break'),
                ('14:45 - 16:00', 'Traffic Integration & Road Scenarios'),
                ('16:00 - 16:45', 'Post-Test Assessment & Performance Review'),
                ('16:45 - 17:00', 'Certificate Presentation & Program Closure')
            ]
        
        agenda_table_day2 = doc.add_table(rows=len(day2_items)+1, cols=2)
        agenda_table_day2.style = 'Light Grid Accent 1'
        agenda_table_day2.rows[0].cells[0].text = 'Time'
        agenda_table_day2.rows[0].cells[1].text = 'Activity'
        for idx, (time, activity) in enumerate(day2_items, 1):
            agenda_table_day2.rows[idx].cells[0].text = time
            agenda_table_day2.rows[idx].cells[1].text = activity
        
        doc.add_page_break()
        
        # TRAINING DETAILS
        doc.add_heading('4. TRAINING DETAILS', 1)
        doc.add_paragraph(f"Program: {program.get('name', 'N/A')}")
        doc.add_paragraph(f"Location: {session.get('location', 'N/A')}")
        doc.add_paragraph(f"Dates: {session.get('start_date', 'N/A')} to {session.get('end_date', 'N/A')}")
        doc.add_paragraph(f"Total Participants: {len(participants)}")
        doc.add_paragraph()
        doc.add_paragraph("Participants List:")
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'ID Number'
        for p in participants:
            row_cells = table.add_row().cells
            row_cells[0].text = p['name']
            row_cells[1].text = str(p['id_number'])
        doc.add_page_break()
        
        # PRE-POST EVALUATION SUMMARY
        doc.add_heading('5. PRE-POST EVALUATION SUMMARY', 1)
        # Summary statistics
        doc.add_paragraph(f"Pre-Test Pass Rate: {pre_pass_count}/{len(participants)} participants ({(pre_pass_count/len(participants)*100):.0f}%)")
        doc.add_paragraph(f"Post-Test Pass Rate: {post_pass_count}/{len(participants)} participants ({(post_pass_count/len(participants)*100):.0f}%)")
        doc.add_paragraph(f"Participants Showing Improvement: {improved_count}/{len(participants)} ({(improved_count/len(participants)*100):.0f}%)")
        doc.add_paragraph(f"Average Score Change: {improvement:+.1f}%")
        doc.add_paragraph()
        
        # Performance Summary Table
        doc.add_paragraph("TABULATED RESULTS:", style='Heading 3')
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Participant'
        hdr_cells[1].text = 'ID Number'
        hdr_cells[2].text = 'Pre-Test'
        hdr_cells[3].text = 'Post-Test'
        hdr_cells[4].text = 'Improvement'
        hdr_cells[5].text = 'Status'
        
        for p in participants:
            row_cells = table.add_row().cells
            row_cells[0].text = p['name']
            row_cells[1].text = str(p['id_number'])
            row_cells[2].text = f"{p['pre_test_score']:.0f}%"
            row_cells[3].text = f"{p['post_test_score']:.0f}%"
            row_cells[4].text = f"{p['improvement']:+.0f}%"
            row_cells[5].text = 'PASS' if p['post_test_passed'] else 'FAIL'
        
        doc.add_page_break()
        
        # DETAILED PERFORMANCE ANALYSIS WITH INSIGHTS
        doc.add_heading('6. DETAILED PERFORMANCE ANALYSIS', 1)
        doc.add_paragraph("Individual participant performance with remarks and recommendations:")
        doc.add_paragraph()
        
        for idx, p in enumerate(participants, 1):
            doc.add_paragraph(f"{idx}. {p['name']} (ID: {p['id_number']})", style='Heading 3')
            perf_text = f"   Pre-Test: {p['pre_test_score']:.0f}% | Post-Test: {p['post_test_score']:.0f}% | Change: {p['improvement']:+.0f}%"
            doc.add_paragraph(perf_text)
            
            # Generate performance remarks based on improvement
            if p['improvement'] >= 20:
                remark = "EXCELLENT IMPROVEMENT - Participant demonstrated exceptional learning and engagement. Strong grasp of defensive techniques."
            elif p['improvement'] >= 10:
                remark = "GOOD IMPROVEMENT - Participant showed solid progress and understanding of safety principles."
            elif p['improvement'] >= 0:
                remark = "SATISFACTORY PROGRESS - Participant maintained or slightly improved performance. Continue practicing learned techniques."
            elif p['improvement'] >= -10:
                remark = "NEEDS ATTENTION - Minor score decrease observed. Recommend follow-up coaching and review of key concepts."
            else:
                remark = "REQUIRES IMMEDIATE SUPPORT - Significant score decrease. Recommend one-on-one coaching session and practical refresher."
            
            # Pass/Fail status remark
            if not p['pre_test_passed'] and p['post_test_passed']:
                remark += " Successfully progressed from FAIL to PASS status."
            elif not p['post_test_passed']:
                remark += " Did not achieve passing score - recommend additional training."
            
            doc.add_paragraph(f"   Remark: {remark}")
            doc.add_paragraph()
        
        # Overall Performance Insights
        doc.add_paragraph("OVERALL PERFORMANCE INSIGHTS:", style='Heading 3')
        high_performers = [p for p in participants if p['improvement'] >= 15]
        needs_support = [p for p in participants if p['improvement'] < 0]
        
        insights = []
        if high_performers:
            insights.append(f"• {len(high_performers)} participant(s) demonstrated excellent improvement (≥15% gain), indicating strong training absorption.")
        if needs_support:
            insights.append(f"• {len(needs_support)} participant(s) showed score decrease and require targeted follow-up support.")
        insights.append(f"• Average improvement of {improvement:+.1f}% indicates {'effective' if improvement > 5 else 'moderate'} training impact.")
        insights.append(f"• Post-test pass rate of {(post_pass_count/len(participants)*100):.0f}% {'meets' if post_pass_count/len(participants) >= 0.8 else 'is below'} target standards.")
        
        for insight in insights:
            doc.add_paragraph(insight)
        
        doc.add_page_break()
        
        # Get chief trainer feedback before displaying
        chief_trainer_feedback = await db.chief_trainer_feedback.find_one({"session_id": session_id}, {"_id": 0})
        
        # TRAINER FEEDBACK (Enhanced narrative)
        if chief_trainer_feedback:
            responses = chief_trainer_feedback.get('responses', {})
            template = await db.feedback_templates.find_one({"id": "chief_trainer_feedback_template"}, {"_id": 0})
            
            doc.add_paragraph(
                "The chief trainer provided comprehensive feedback on the training delivery, participant engagement, "
                "and safety observations throughout the program. Key observations and recommendations are detailed below:"
            )
            doc.add_paragraph()
            
            # Extract narrative responses from chief trainer
            for question_id, answer in responses.items():
                if template:
                    for q in template.get('questions', []):
                        if q.get('id') == question_id:
                            doc.add_paragraph(f"{q.get('question')}:", style='Heading 3')
                            if q.get('type') == 'rating':
                                stars = '⭐' * int(answer) if isinstance(answer, (int, float)) else answer
                                doc.add_paragraph(f"   Rating: {stars} ({answer}/{q.get('scale', 5)})")
                            else:
                                doc.add_paragraph(f"   {answer}")
                            doc.add_paragraph()
            
            # Add professional summary quote
            doc.add_paragraph()
            doc.add_paragraph(
                "The trainer observed that participants were highly engaged and receptive to feedback. "
                "Safety issues identified during vehicle inspections were communicated to participants and management. "
                "Overall, the training environment was conducive to learning with participants demonstrating strong "
                "commitment to improving their safety practices."
            )
        else:
            doc.add_paragraph("[Chief Trainer feedback pending submission]")
        
        doc.add_page_break()
        
        # TRAINING PHOTOS
        doc.add_heading('8. TRAINING PHOTOS', 1)
        if training_photos['group_photo']:
            doc.add_paragraph("Group Photo:", style='Heading 3')
            doc.add_paragraph(f"[Photo URL: {training_photos['group_photo']}]")
            doc.add_paragraph()
        
        if training_photos['theory_photo_1'] or training_photos['theory_photo_2']:
            doc.add_paragraph("Theory Session Photos:", style='Heading 3')
            if training_photos['theory_photo_1']:
                doc.add_paragraph(f"[Photo 1 URL: {training_photos['theory_photo_1']}]")
            if training_photos['theory_photo_2']:
                doc.add_paragraph(f"[Photo 2 URL: {training_photos['theory_photo_2']}]")
            doc.add_paragraph()
        
        if training_photos['practical_photo_1'] or training_photos['practical_photo_2'] or training_photos['practical_photo_3']:
            doc.add_paragraph("Practical Session Photos:", style='Heading 3')
            if training_photos['practical_photo_1']:
                doc.add_paragraph(f"[Photo 1 URL: {training_photos['practical_photo_1']}]")
            if training_photos['practical_photo_2']:
                doc.add_paragraph(f"[Photo 2 URL: {training_photos['practical_photo_2']}]")
            if training_photos['practical_photo_3']:
                doc.add_paragraph(f"[Photo 3 URL: {training_photos['practical_photo_3']}]")
        
        doc.add_page_break()
        
        # PARTICIPANT FEEDBACK SUMMARY (Enhanced)
        doc.add_heading('9. PARTICIPANT FEEDBACK SUMMARY', 1)
        if feedback_data:
            # Calculate average star ratings
            star_questions = []
            text_questions = []
            
            # Categorize questions
            if feedback_data:
                for response in feedback_data[0]['responses']:
                    if isinstance(response['answer'], int):
                        star_questions.append(response['question'])
                    else:
                        text_questions.append(response['question'])
            
            # PART 1: QUANTITATIVE FEEDBACK
            if star_questions:
                doc.add_paragraph("A. QUANTITATIVE FEEDBACK (Rating Scores):", style='Heading 3')
                doc.add_paragraph("Average ratings across all participants on a 5-point scale:")
                doc.add_paragraph()
                
                for question in star_questions:
                    ratings = [r['answer'] for fb in feedback_data for r in fb['responses'] if r['question'] == question and isinstance(r['answer'], int)]
                    if ratings:
                        avg_rating = sum(ratings) / len(ratings)
                        stars = '⭐' * int(round(avg_rating))
                        doc.add_paragraph(f"• {question}: {stars} ({avg_rating:.1f}/5.0)")
                
                # Overall satisfaction calculation
                all_ratings = [r['answer'] for fb in feedback_data for r in fb['responses'] if isinstance(r['answer'], int)]
                if all_ratings:
                    overall_avg = sum(all_ratings) / len(all_ratings)
                    doc.add_paragraph()
                    doc.add_paragraph(f"OVERALL SATISFACTION: {'⭐' * int(round(overall_avg))} ({overall_avg:.1f}/5.0)", style='Heading 3')
                doc.add_paragraph()
            
            # PART 2: QUALITATIVE FEEDBACK THEMES
            if text_questions:
                doc.add_paragraph("B. QUALITATIVE FEEDBACK (Key Themes):", style='Heading 3')
                
                # Collect all text responses
                all_text_responses = []
                for fb in feedback_data:
                    for response in fb['responses']:
                        if not isinstance(response['answer'], int):
                            all_text_responses.append(response['answer'])
                
                # Analyze common themes (simple keyword matching)
                positive_keywords = ['good', 'excellent', 'great', 'helpful', 'informative', 'clear', 'effective']
                improvement_keywords = ['more', 'extend', 'longer', 'additional', 'better', 'improve']
                
                positive_count = sum(1 for resp in all_text_responses if any(kw in str(resp).lower() for kw in positive_keywords))
                improvement_count = sum(1 for resp in all_text_responses if any(kw in str(resp).lower() for kw in improvement_keywords))
                
                doc.add_paragraph(f"• Positive Remarks: {positive_count} participants expressed satisfaction with training delivery and content")
                if improvement_count > 0:
                    doc.add_paragraph(f"• Improvement Suggestions: {improvement_count} participants suggested enhancements (e.g., extended duration, additional videos)")
                doc.add_paragraph()
                
                # PART 3: INDIVIDUAL RESPONSES (Detailed)
                doc.add_paragraph("C. DETAILED INDIVIDUAL RESPONSES:", style='Heading 3')
                for idx, fb in enumerate(feedback_data, 1):
                    doc.add_paragraph(f"{idx}. {fb['participant_name']}", style='Heading 4')
                    for response in fb['responses']:
                        if not isinstance(response['answer'], int):  # Text responses
                            doc.add_paragraph(f"   Q: {response['question']}")
                            doc.add_paragraph(f"   A: {response['answer']}")
                            doc.add_paragraph()
        else:
            doc.add_paragraph("No feedback submitted yet.")
        
        doc.add_page_break()
        
        # MOTORCYCLE/VEHICLE CONDITION & EMPLOYER RECOMMENDATIONS (Enhanced)
        doc.add_heading('10. VEHICLE CONDITION ASSESSMENT & EMPLOYER RECOMMENDATIONS', 1)
        
        if vehicle_issues:
            doc.add_paragraph(
                f"During the training program, pre-ride safety inspections were conducted on all participant "
                f"{'motorcycles' if is_motorcycle else 'vehicles'}. The inspections revealed {len(vehicle_issues)} "
                f"{'motorcycles' if is_motorcycle else 'vehicles'} with safety concerns that require immediate attention."
            )
            doc.add_paragraph()
            
            # PART A: SAFETY ISSUES IDENTIFIED
            doc.add_paragraph("A. SAFETY ISSUES IDENTIFIED:", style='Heading 3')
            for vehicle_issue in vehicle_issues:
                doc.add_paragraph(f"Participant: {vehicle_issue['participant_name']}", style='Heading 4')
                for issue in vehicle_issue['issues']:
                    doc.add_paragraph(f"   • {issue['item']}: {issue['comment']}")
                    if issue['photo_url']:
                        doc.add_paragraph(f"     [Photo Evidence: {issue['photo_url']}]")
                doc.add_paragraph()
            
            # PART B: SAFETY IMPLICATIONS
            doc.add_paragraph("B. SAFETY IMPLICATIONS:", style='Heading 3')
            common_issues = {}
            for vehicle_issue in vehicle_issues:
                for issue in vehicle_issue['issues']:
                    item_category = issue['item'].lower()
                    if 'tyre' in item_category or 'tire' in item_category:
                        common_issues['worn_tyres'] = common_issues.get('worn_tyres', 0) + 1
                    elif 'lamp' in item_category or 'light' in item_category:
                        common_issues['faulty_lamps'] = common_issues.get('faulty_lamps', 0) + 1
                    elif 'chain' in item_category:
                        common_issues['loose_chains'] = common_issues.get('loose_chains', 0) + 1
                    elif 'mirror' in item_category:
                        common_issues['missing_mirrors'] = common_issues.get('missing_mirrors', 0) + 1
                    elif 'ppe' in item_category or 'helmet' in item_category or 'jacket' in item_category:
                        common_issues['ppe_issues'] = common_issues.get('ppe_issues', 0) + 1
            
            if common_issues:
                for issue_type, count in common_issues.items():
                    if issue_type == 'worn_tyres':
                        doc.add_paragraph(f"• Worn Tyres ({count} cases): Increased risk of skidding and loss of control, especially in wet conditions")
                    elif issue_type == 'faulty_lamps':
                        doc.add_paragraph(f"• Faulty Lamps/Lights ({count} cases): Reduced visibility at night, increased accident risk")
                    elif issue_type == 'loose_chains':
                        doc.add_paragraph(f"• Loose Chains ({count} cases): Risk of chain breakage leading to loss of control")
                    elif issue_type == 'missing_mirrors':
                        doc.add_paragraph(f"• Missing/Damaged Mirrors ({count} cases): Impaired situational awareness and blind spot monitoring")
                    elif issue_type == 'ppe_issues':
                        doc.add_paragraph(f"• PPE Non-Compliance ({count} cases): Increased severity of injuries in case of accidents")
            doc.add_paragraph()
            
            # PART C: RECOMMENDATIONS FOR EMPLOYER
            doc.add_paragraph("C. RECOMMENDATIONS FOR EMPLOYER:", style='Heading 3')
            recommendations = [
                "1. IMMEDIATE ACTION REQUIRED:",
                f"   • Conduct immediate safety inspections on all {len(vehicle_issues)} flagged {'motorcycles' if is_motorcycle else 'vehicles'}",
                "   • Ground vehicles until critical safety issues are resolved",
                "   • Provide temporary alternative transportation if needed",
                "",
                "2. ESTABLISH REGULAR MAINTENANCE PROTOCOL:",
                f"   • Implement monthly pre-ride safety inspection checklist for all {'motorcycles' if is_motorcycle else 'vehicles'}",
                "   • Assign designated personnel for routine maintenance verification",
                "   • Maintain detailed maintenance logs for each vehicle",
                "",
                "3. PPE COMPLIANCE:",
                "   • Enforce mandatory PPE usage policy (helmet, jacket, gloves, boots)",
                "   • Provide company-issued PPE if necessary",
                "   • Conduct regular PPE condition checks",
                "",
                "4. INTEGRATE INTO SAFETY SOP:",
                "   • Include vehicle inspection as part of daily work routine",
                "   • Establish clear reporting channels for safety issues",
                "   • Implement consequences for non-compliance",
                "",
                "5. SUPPLEMENTARY TRAINING:",
                "   • Conduct basic vehicle maintenance workshop for employees",
                "   • Provide refresher training on pre-ride safety checks"
            ]
            for rec in recommendations:
                doc.add_paragraph(rec)
        else:
            doc.add_paragraph("✓ EXCELLENT RESULT: All vehicles inspected were found to be in good working condition with no safety concerns identified.")
            doc.add_paragraph()
            doc.add_paragraph(
                "This indicates strong commitment to vehicle maintenance and safety standards. We recommend "
                "continuing current maintenance practices and conducting regular quarterly safety inspections."
            )
        
        doc.add_page_break()
        
        # COORDINATOR FEEDBACK (Enhanced)
        doc.add_heading('11. COORDINATOR FEEDBACK', 1)
        coordinator_feedback = await db.coordinator_feedback.find_one({"session_id": session_id}, {"_id": 0})
        if coordinator_feedback:
            doc.add_paragraph(
                "The training coordinator provided comprehensive observations on logistics, participant engagement, "
                "and overall program execution. Key observations and recommendations are detailed below:"
            )
            doc.add_paragraph()
            
            responses = coordinator_feedback.get('responses', {})
            for question_id, answer in responses.items():
                # Get question text from template
                template = await db.feedback_templates.find_one({"id": "coordinator_feedback_template"}, {"_id": 0})
                if template:
                    for q in template.get('questions', []):
                        if q.get('id') == question_id:
                            doc.add_paragraph(f"{q.get('question')}:", style='Heading 3')
                            if q.get('type') == 'rating':
                                stars = '⭐' * int(answer) if isinstance(answer, (int, float)) else answer
                                doc.add_paragraph(f"   Rating: {stars} ({answer}/{q.get('scale', 5)})")
                            else:
                                doc.add_paragraph(f"   {answer}")
                            doc.add_paragraph()
            
            # Add formal closing
            doc.add_paragraph()
            doc.add_paragraph(
                f"The coordinator acknowledges the strong collaboration between {company.get('name', 'the company')}, "
                "MDDRC training team, and participants throughout the program. Participants demonstrated excellent "
                "discipline and commitment to learning, contributing to the overall success of the training initiative."
            )
        else:
            doc.add_paragraph("[Coordinator feedback pending submission]")
        
        doc.add_page_break()
        
        # RECOMMENDATIONS MOVING FORWARD
        doc.add_heading('12. RECOMMENDATIONS MOVING FORWARD', 1)
        doc.add_paragraph(
            "Based on the training outcomes, participant feedback, and safety observations, "
            "the following recommendations are proposed to sustain and enhance the safety culture:"
        )
        doc.add_paragraph()
        
        recommendations_forward = [
            "1. ENFORCE PRE-RIDE/PRE-DRIVE SAFETY CHECKS:",
            f"   • Mandate daily pre-{'ride' if is_motorcycle else 'drive'} safety inspections using a standardized checklist",
            "   • Implement digital logging system for inspection records",
            "   • Designate safety officers to conduct random spot checks",
            "",
            "2. MONTHLY VERIFICATION PROGRAM:",
            "   • Conduct monthly vehicle condition audits",
            "   • Schedule preventive maintenance based on mileage/usage",
            "   • Track and analyze vehicle-related incidents",
            "",
            "3. MAINTENANCE SUPPORT:",
            "   • Establish partnerships with authorized service centers for employee discounts",
            "   • Provide maintenance subsidy program for safety-critical components",
            "   • Create emergency maintenance fund for immediate safety repairs",
            "",
            "4. POST-TRAINING MATERIALS:",
            "   • Distribute safety reminder cards or posters for display",
            "   • Share digital safety tips via company communication channels",
            "   • Conduct quarterly safety awareness campaigns",
            "",
            "5. TAILOR PRACTICALS TO CLIENT ROUTES:",
            "   • Identify high-risk routes and areas commonly used by employees",
            "   • Conduct route-specific safety briefings",
            "   • Share incident hotspot maps and avoidance strategies",
            "",
            "6. PROMOTE SAFETY CULTURE:",
            "   • Recognize and reward safe riding/driving behavior",
            "   • Establish peer mentorship program for new employees",
            "   • Include safety KPIs in performance evaluations",
            "",
            "7. FOLLOW-UP FOR OUTLIERS:",
            f"   • Provide one-on-one coaching for {len([p for p in participants if p['improvement'] < 0])} participants who showed score decrease" if any(p['improvement'] < 0 for p in participants) else "   • Continue monitoring participant performance in real-world scenarios",
            "   • Conduct 3-month post-training assessment to measure retention",
            "   • Offer refresher training for employees showing concerning behavior"
        ]
        
        for rec in recommendations_forward:
            doc.add_paragraph(rec)
        
        doc.add_page_break()
        
        # CONCLUSION
        doc.add_heading('13. CONCLUSION', 1)
        doc.add_paragraph(
            f"The Defensive {'Riding' if is_motorcycle else 'Driving'} Training conducted for "
            f"{company.get('name', 'N/A')} from {session.get('start_date', 'N/A')} to {session.get('end_date', 'N/A')} "
            f"was successfully completed with {len(participants)} participants demonstrating measurable improvement in "
            f"safety awareness and defensive {'riding' if is_motorcycle else 'driving'} competencies."
        )
        doc.add_paragraph()
        
        doc.add_paragraph(
            f"Key achievements include an average score improvement of {improvement:+.1f}%, "
            f"a post-training pass rate of {(post_pass_count/len(participants)*100):.0f}%, and high participant "
            f"satisfaction levels. The training successfully enhanced hazard recognition skills, emergency response "
            f"techniques, and safety-first mindset among participants."
        )
        doc.add_paragraph()
        
        if vehicle_issues:
            doc.add_paragraph(
                f"Vehicle safety inspections identified {len(vehicle_issues)} {'motorcycles' if is_motorcycle else 'vehicles'} "
                "requiring immediate attention. Detailed recommendations have been provided to address these concerns "
                "and prevent potential accidents."
            )
            doc.add_paragraph()
        
        doc.add_paragraph(
            "MDDRC extends sincere appreciation to the management and employees of "
            f"{company.get('name', 'the company')} for their strong collaboration and commitment throughout this program. "
            "The enthusiastic participation and positive learning attitude demonstrated by all participants contributed "
            "significantly to the program's success."
        )
        doc.add_paragraph()
        
        doc.add_paragraph(
            "We remain committed to supporting your organization's journey towards a safer workplace and look forward "
            "to continued partnership in promoting road safety excellence."
        )
        
        doc.add_page_break()
        
        # APPENDICES
        doc.add_heading('APPENDICES', 1)
        
        # APPENDIX A: Pre & Post Test Raw Scores
        doc.add_heading('Appendix A: Pre & Post Test Raw Scores', 2)
        appendix_table = doc.add_table(rows=len(participants)+1, cols=5)
        appendix_table.style = 'Light Grid Accent 1'
        hdr = appendix_table.rows[0].cells
        hdr[0].text = 'No.'
        hdr[1].text = 'Participant Name'
        hdr[2].text = 'Pre-Test Score'
        hdr[3].text = 'Post-Test Score'
        hdr[4].text = 'Improvement'
        
        for idx, p in enumerate(participants, 1):
            row = appendix_table.rows[idx].cells
            row[0].text = str(idx)
            row[1].text = p['name']
            row[2].text = f"{p['pre_test_score']:.0f}%"
            row[3].text = f"{p['post_test_score']:.0f}%"
            row[4].text = f"{p['improvement']:+.0f}%"
        
        doc.add_page_break()
        
        # APPENDIX B: Vehicle Condition Photos
        if vehicle_issues:
            doc.add_heading('Appendix B: Vehicle Condition Photos', 2)
            doc.add_paragraph("Photographic evidence of safety issues identified during vehicle inspections:")
            doc.add_paragraph()
            for vehicle_issue in vehicle_issues:
                doc.add_paragraph(f"{vehicle_issue['participant_name']}:", style='Heading 4')
                for issue in vehicle_issue['issues']:
                    if issue['photo_url']:
                        doc.add_paragraph(f"• {issue['item']}")
                        doc.add_paragraph(f"  [Photo URL: {issue['photo_url']}]")
                doc.add_paragraph()
            doc.add_page_break()
        
        # APPENDIX C: Feedback Form Summary
        doc.add_heading('Appendix C: Participant Feedback Form Summary', 2)
        if feedback_data:
            doc.add_paragraph("Complete participant feedback responses:")
            doc.add_paragraph()
            for idx, fb in enumerate(feedback_data, 1):
                doc.add_paragraph(f"{idx}. {fb['participant_name']}", style='Heading 4')
                for response in fb['responses']:
                    doc.add_paragraph(f"   Q: {response['question']}")
                    doc.add_paragraph(f"   A: {response['answer']}")
                doc.add_paragraph()
        else:
            doc.add_paragraph("[No feedback data available]")
        
        doc.add_page_break()
        
        # SIGNATURES
        doc.add_heading('APPROVAL & SIGNATURES', 1)
        doc.add_paragraph()
        sig_table = doc.add_table(rows=4, cols=2)
        sig_table.style = 'Light List'
        
        sig_table.rows[0].cells[0].text = 'Prepared by:'
        sig_table.rows[0].cells[1].text = ''
        sig_table.rows[1].cells[0].text = 'Name:'
        sig_table.rows[1].cells[1].text = current_user.full_name
        sig_table.rows[2].cells[0].text = 'Position:'
        sig_table.rows[2].cells[1].text = 'Training Coordinator'
        sig_table.rows[3].cells[0].text = 'Date:'
        sig_table.rows[3].cells[1].text = get_malaysia_time().strftime('%Y-%m-%d')
        
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("_" * 60)
        doc.add_paragraph()
        
        sig_table2 = doc.add_table(rows=4, cols=2)
        sig_table2.style = 'Light List'
        sig_table2.rows[0].cells[0].text = 'Reviewed & Approved by:'
        sig_table2.rows[0].cells[1].text = ''
        sig_table2.rows[1].cells[0].text = 'Name:'
        sig_table2.rows[1].cells[1].text = '________________________'
        sig_table2.rows[2].cells[0].text = 'Position:'
        sig_table2.rows[2].cells[1].text = 'Person In Charge / Supervisor'
        sig_table2.rows[3].cells[0].text = 'Date:'
        sig_table2.rows[3].cells[1].text = '________________________'
        
        doc.add_paragraph()
        doc.add_paragraph()
        footer = doc.add_paragraph('--- END OF REPORT ---')
        footer.alignment = 1
        
        doc.add_page_break()
        
        # SIGNATURES
        doc.add_heading('11. SIGNATURES', 1)
        doc.add_paragraph()
        doc.add_paragraph("_" * 40)
        doc.add_paragraph(f"Coordinator: {current_user.full_name}")
        doc.add_paragraph(f"Date: ________________")
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("_" * 40)
        doc.add_paragraph("PIC/Supervisor Signature")
        doc.add_paragraph(f"Date: ________________")
        
        # Save DOCX
        report_filename = f"Training_Report_{session_id}_{get_malaysia_time().strftime('%Y%m%d_%H%M%S')}.docx"
        report_path = REPORT_DIR / report_filename
        doc.save(str(report_path))
        
        # Update training report record with DOCX filename
        await db.training_reports.update_one(
            {"session_id": session_id},
            {"$set": {"docx_filename": report_filename, "generated_at": get_malaysia_time().isoformat()}},
            upsert=True
        )
        
        return {
            "message": "DOCX report generated successfully",
            "filename": report_filename,
            "download_url": f"/api/training-reports/{session_id}/download-docx"
        }
        
    except Exception as e:
        logging.error(f"Failed to generate DOCX report: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate report: {str(e)}")

@api_router.get("/training-reports/{session_id}/download-docx")
async def download_docx_report(session_id: str, current_user: User = Depends(get_current_user)):
    """Download the generated DOCX report"""
    
    if current_user.role not in ["coordinator", "admin", "supervisor"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    # Get report filename from database
    training_report = await db.training_reports.find_one({"session_id": session_id}, {"_id": 0})
    
    if not training_report or not training_report.get('docx_filename'):
        raise HTTPException(status_code=404, detail="Report not found. Please generate it first.")
    
    report_path = REPORT_DIR / training_report['docx_filename']
    
    if not report_path.exists():
        raise HTTPException(status_code=404, detail="Report file not found")
    
    return FileResponse(
        path=str(report_path),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=training_report['docx_filename'],
        headers={"Content-Disposition": f"attachment; filename={training_report['docx_filename']}"}
    )

@api_router.post("/training-reports/{session_id}/upload-edited-docx")
async def upload_edited_docx(
    session_id: str,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    """Upload edited DOCX report"""
    
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Only coordinators and admins can upload reports")
    
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only DOCX files are allowed")
    
    try:
        # Save edited DOCX
        edited_filename = f"Training_Report_{session_id}_edited_{get_malaysia_time().strftime('%Y%m%d_%H%M%S')}.docx"
        edited_path = REPORT_DIR / edited_filename
        
        with open(edited_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Update database
        await db.training_reports.update_one(
            {"session_id": session_id},
            {"$set": {
                "edited_docx_filename": edited_filename,
                "uploaded_at": get_malaysia_time().isoformat()
            }},
            upsert=True
        )
        
        return {
            "message": "Edited report uploaded successfully",
            "filename": edited_filename
        }
        
    except Exception as e:
        logging.error(f"Failed to upload edited report: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to upload report: {str(e)}")



@api_router.post("/training-reports/{session_id}/upload-final-pdf")
async def upload_final_pdf_report(
    session_id: str,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    """Upload final edited PDF report"""
    
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Only coordinators and admins can upload reports")
    
    # Validate file type
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are accepted")
    
    # Check file size (max 20MB)
    file.file.seek(0, 2)
    file_size = file.file.tell()
    file.file.seek(0)
    
    max_size = 20 * 1024 * 1024  # 20MB
    if file_size > max_size:
        raise HTTPException(status_code=400, detail="File size exceeds 20MB limit")
    
    try:
        # Save final PDF
        pdf_filename = f"Training_Report_{session_id}_final_{get_malaysia_time().strftime('%Y%m%d_%H%M%S')}.pdf"
        pdf_path = REPORT_PDF_DIR / pdf_filename
        
        with open(pdf_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Get session and program details
        session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
        program = None
        company = None
        if session:
            if session.get('program_id'):
                program = await db.programs.find_one({"id": session['program_id']}, {"_id": 0})
            if session.get('company_id'):
                company = await db.companies.find_one({"id": session['company_id']}, {"_id": 0})
        
        # Update database with submitted status
        await db.training_reports.update_one(
            {"session_id": session_id},
            {"$set": {
                "final_pdf_filename": pdf_filename,
                "pdf_url": f"/api/static/reports_pdf/{pdf_filename}",
                "status": "submitted",
                "submitted_at": get_malaysia_time().isoformat(),
                "submitted_by": current_user.id,
                "program_id": program.get('id') if program else None,
                "company_id": company.get('id') if company else None,
                "session_name": session.get('name') if session else None,
                "session_start_date": session.get('start_date') if session else None,
                "session_end_date": session.get('end_date') if session else None
            }},
            upsert=True
        )
        
        return {
            "message": "Final report submitted successfully",
            "filename": pdf_filename,
            "pdf_url": f"/api/static/reports_pdf/{pdf_filename}"
        }
        
    except Exception as e:
        logging.error(f"Failed to upload final PDF: {str(e)}")


# Get submitted reports for supervisor
@api_router.get("/training-reports/supervisor/sessions")
async def get_supervisor_reports(current_user: User = Depends(get_current_user)):
    """Get all submitted reports for sessions assigned to supervisor"""
    
    if current_user.role not in ["supervisor", "admin"]:
        raise HTTPException(status_code=403, detail="Only supervisors and admins can access this")
    
    # Get sessions assigned to supervisor
    if current_user.role == "supervisor":
        sessions = await db.sessions.find({"supervisor_id": current_user.id}, {"_id": 0}).to_list(100)
    else:
        # Admin can see all
        sessions = await db.sessions.find({}, {"_id": 0}).to_list(1000)
    
    session_ids = [s['id'] for s in sessions]
    
    # Get submitted reports for these sessions
    reports = await db.training_reports.find(
        {
            "session_id": {"$in": session_ids},
            "status": "submitted"
        },
        {"_id": 0}
    ).to_list(100)
    
    # Enrich with session details
    enriched_reports = []
    for report in reports:
        session = next((s for s in sessions if s['id'] == report['session_id']), None)
        if session:
            enriched_reports.append({
                **report,
                "session_name": session.get('name'),
                "session_start_date": session.get('start_date'),
                "session_end_date": session.get('end_date'),
                "location": session.get('location')
            })
    
    return enriched_reports

@api_router.post("/training-reports/{session_id}/submit-final")
async def submit_final_report(session_id: str, current_user: User = Depends(get_current_user)):
    """Submit final report - converts to PDF and notifies supervisor/admin"""
    
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Only coordinators and admins can submit reports")
    
    try:
        # Get the latest report (edited if exists, otherwise generated)
        training_report = await db.training_reports.find_one({"session_id": session_id}, {"_id": 0})
        
        if not training_report:
            raise HTTPException(status_code=404, detail="No report found. Please generate a report first.")
        
        docx_filename = training_report.get('edited_docx_filename') or training_report.get('docx_filename')
        
        if not docx_filename:
            raise HTTPException(status_code=404, detail="No report file found")
        
        docx_path = REPORT_DIR / docx_filename
        
        if not docx_path.exists():
            raise HTTPException(status_code=404, detail="Report file not found")
        
        # Convert DOCX to PDF using LibreOffice
        pdf_filename = docx_filename.replace('.docx', '.pdf')
        pdf_path = REPORT_PDF_DIR / pdf_filename
        
        subprocess.run([
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(REPORT_PDF_DIR),
            str(docx_path)
        ], check=True)
        
        # Update training report status
        await db.training_reports.update_one(
            {"session_id": session_id},
            {"$set": {
                "pdf_filename": pdf_filename,
                "status": "submitted",
                "submitted_at": get_malaysia_time().isoformat(),
                "submitted_by": current_user.id
            }}
        )
        
        # Get session and create notifications for supervisor and admin
        session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
        
        # Notify supervisor
        if session.get('supervisor_ids'):
            for supervisor_id in session['supervisor_ids']:
                await db.notifications.insert_one({
                    "id": str(uuid.uuid4()),
                    "user_id": supervisor_id,
                    "type": "training_report_submitted",
                    "message": f"Training report for {session.get('name')} has been submitted",
                    "session_id": session_id,
                    "read": False,
                    "created_at": get_malaysia_time().isoformat()
                })
        
        # Notify all admins
        admins = await db.users.find({"role": "admin"}, {"_id": 0}).to_list(100)
        for admin in admins:
            await db.notifications.insert_one({
                "id": str(uuid.uuid4()),
                "user_id": admin['id'],
                "type": "training_report_submitted",
                "message": f"Training report for {session.get('name')} has been submitted by {current_user.full_name}",
                "session_id": session_id,
                "read": False,
                "created_at": get_malaysia_time().isoformat()
            })
        
        return {
            "message": "Report submitted successfully and PDF generated",
            "pdf_filename": pdf_filename,
            "download_url": f"/api/training-reports/{session_id}/download-pdf"
        }
        
    except subprocess.CalledProcessError as e:
        logging.error(f"PDF conversion failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to convert report to PDF")
    except Exception as e:
        logging.error(f"Failed to submit report: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to submit report: {str(e)}")

@api_router.get("/training-reports/{session_id}/download-pdf")
async def download_pdf_report(session_id: str, current_user: User = Depends(get_current_user)):
    """Download the final PDF report"""
    
    if current_user.role not in ["coordinator", "admin", "supervisor"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    # Get report filename from database
    training_report = await db.training_reports.find_one({"session_id": session_id}, {"_id": 0})
    
    if not training_report or not training_report.get('pdf_filename'):
        raise HTTPException(status_code=404, detail="PDF report not found. Please submit the report first.")
    
    pdf_path = REPORT_PDF_DIR / training_report['pdf_filename']
    
    if not pdf_path.exists():
        raise HTTPException(status_code=404, detail="PDF file not found")
    
    return FileResponse(
        path=str(pdf_path),
        media_type='application/pdf',
        filename=training_report['pdf_filename']
    )

# Trainer Checklist Routes
@api_router.post("/trainer-checklist/submit")
async def submit_trainer_checklist(checklist_data: TrainerChecklistSubmit, current_user: User = Depends(get_current_user)):
    if current_user.role != "trainer":
        raise HTTPException(status_code=403, detail="Only trainers can submit checklists")
    
    # Create checklist
    checklist_obj = VehicleChecklist(
        participant_id=checklist_data.participant_id,
        session_id=checklist_data.session_id,
        interval="trainer_inspection",
        checklist_items=[item.model_dump() for item in checklist_data.items],
        verified_by=current_user.id,
        verified_at=datetime.now(timezone.utc),
        verification_status="completed"
    )
    
    doc = checklist_obj.model_dump()
    doc['submitted_at'] = doc['submitted_at'].isoformat()
    doc['verified_at'] = doc['verified_at'].isoformat()
    
    await db.vehicle_checklists.insert_one(doc)
    
    # If chief trainer submitted comments, save to session
    if checklist_data.chief_trainer_comments:
        session = await db.sessions.find_one({"id": checklist_data.session_id}, {"_id": 0})
        if session:
            # Check if current trainer is chief
            trainer_assignments = session.get('trainer_assignments', [])
            is_chief = any(t['trainer_id'] == current_user.id and t.get('role') == 'chief' for t in trainer_assignments)
            
            if is_chief:
                await db.sessions.update_one(
                    {"id": checklist_data.session_id},
                    {"$set": {
                        "chief_trainer_comments": checklist_data.chief_trainer_comments,
                        "chief_trainer_id": current_user.id,
                        "chief_trainer_name": current_user.full_name,
                        "comments_submitted_at": get_malaysia_time().isoformat()
                    }}
                )
    
    return {"message": "Checklist submitted successfully", "checklist_id": checklist_obj.id}

@api_router.get("/trainer-checklist/{session_id}/assigned-participants")
async def get_assigned_participants(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "trainer":
        raise HTTPException(status_code=403, detail="Only trainers can access this")
    
    # Get session
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Get all trainers in session
    trainer_assignments = session.get('trainer_assignments', [])
    trainers = [t['trainer_id'] for t in trainer_assignments]
    
    if not trainers:
        return []
    
    # Auto-assign participants to trainers
    participant_ids = session.get('participant_ids', [])
    total_participants = len(participant_ids)
    total_trainers = len(trainers)
    
    if total_trainers == 0:
        return []
    
    # EQUAL DISTRIBUTION: All trainers (chief and regular) get equal number of participants
    # Divide participants equally among all trainers
    participants_per_trainer = total_participants // total_trainers
    remainder = total_participants % total_trainers
    
    # Find current trainer's index in the list
    current_trainer_index = trainers.index(current_user.id)
    
    # Calculate start index and count for this trainer
    start_index = current_trainer_index * participants_per_trainer
    assigned_count = participants_per_trainer
    
    # Distribute remainder evenly (first N trainers get +1 participant)
    if current_trainer_index < remainder:
        assigned_count += 1
    
    end_index = start_index + assigned_count
    assigned_participant_ids = participant_ids[start_index:end_index]
    
    # Get participant details
    participants = await db.users.find(
        {"id": {"$in": assigned_participant_ids}},
        {"_id": 0, "password": 0}
    ).to_list(100)
    
    # Get vehicle details for each
    for participant in participants:
        vehicle = await db.vehicle_details.find_one({
            "participant_id": participant['id'],
            "session_id": session_id
        }, {"_id": 0})
        participant['vehicle_details'] = vehicle
        
        # Get existing checklist
        checklist = await db.vehicle_checklists.find_one({
            "participant_id": participant['id'],
            "session_id": session_id,
            "verified_by": current_user.id
        }, {"_id": 0})
        participant['checklist'] = checklist
    
    return participants

# Vehicle Checklist Routes
@api_router.post("/checklists/submit", response_model=VehicleChecklist)
async def submit_checklist(checklist_data: ChecklistSubmit, current_user: User = Depends(get_current_user)):
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can submit checklists")
    
    checklist_obj = VehicleChecklist(
        participant_id=current_user.id,
        session_id=checklist_data.session_id,
        interval=checklist_data.interval,
        checklist_items=checklist_data.checklist_items
    )
    
    doc = checklist_obj.model_dump()
    doc['submitted_at'] = doc['submitted_at'].isoformat()
    if doc.get('verified_at'):
        doc['verified_at'] = doc['verified_at'].isoformat()
    
    await db.vehicle_checklists.insert_one(doc)
    
    await db.participant_access.update_one(
        {"participant_id": current_user.id, "session_id": checklist_data.session_id},
        {"$set": {"checklist_submitted": True}}
    )
    
    return checklist_obj

@api_router.get("/checklists/participant/{participant_id}")
async def get_participant_checklists(participant_id: str, current_user: User = Depends(get_current_user)):
    """Get all checklists for a participant (completed by trainers)"""
    # Allow participant themselves, trainers, coordinators, and admins
    if current_user.role not in ["trainer", "coordinator", "admin"] and current_user.id != participant_id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    checklists = await db.vehicle_checklists.find({
        "participant_id": participant_id
    }, {"_id": 0}).to_list(1000)
    
    for checklist in checklists:
        if isinstance(checklist.get('submitted_at'), str):
            checklist['submitted_at'] = datetime.fromisoformat(checklist['submitted_at'])
        if checklist.get('verified_at') and isinstance(checklist['verified_at'], str):
            checklist['verified_at'] = datetime.fromisoformat(checklist['verified_at'])
    
    return checklists

@api_router.get("/vehicle-checklists/{session_id}/{participant_id}")
async def get_checklist(session_id: str, participant_id: str, current_user: User = Depends(get_current_user)):
    # Allow trainer, coordinator, admin, or the participant themselves
    if current_user.role not in ["trainer", "coordinator", "admin"] and current_user.id != participant_id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    checklist = await db.vehicle_checklists.find_one({
        "participant_id": participant_id,
        "session_id": session_id
    }, {"_id": 0})
    
    if not checklist:
        raise HTTPException(status_code=404, detail="Checklist not found")
    
    if isinstance(checklist.get('submitted_at'), str):
        checklist['submitted_at'] = datetime.fromisoformat(checklist['submitted_at'])
    if checklist.get('verified_at') and isinstance(checklist['verified_at'], str):
        checklist['verified_at'] = datetime.fromisoformat(checklist['verified_at'])
    
    return checklist

@api_router.get("/checklists/participant/{participant_id}", response_model=List[VehicleChecklist])
async def get_participant_checklists(participant_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role == "participant" and current_user.id != participant_id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    checklists = await db.vehicle_checklists.find({"participant_id": participant_id}, {"_id": 0}).to_list(100)
    for checklist in checklists:
        if isinstance(checklist.get('submitted_at'), str):
            checklist['submitted_at'] = datetime.fromisoformat(checklist['submitted_at'])
        if checklist.get('verified_at') and isinstance(checklist['verified_at'], str):
            checklist['verified_at'] = datetime.fromisoformat(checklist['verified_at'])
    return checklists

@api_router.get("/checklists/pending", response_model=List[VehicleChecklist])
async def get_pending_checklists(current_user: User = Depends(get_current_user)):
    if current_user.role != "supervisor" and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only supervisors can verify checklists")
    
    checklists = await db.vehicle_checklists.find({"verification_status": "pending"}, {"_id": 0}).to_list(100)
    for checklist in checklists:
        if isinstance(checklist.get('submitted_at'), str):
            checklist['submitted_at'] = datetime.fromisoformat(checklist['submitted_at'])
    return checklists

@api_router.post("/checklists/verify")
async def verify_checklist(verification: ChecklistVerify, current_user: User = Depends(get_current_user)):
    if current_user.role != "supervisor" and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only supervisors can verify checklists")
    
    result = await db.vehicle_checklists.update_one(
        {"id": verification.checklist_id},
        {
            "$set": {
                "verification_status": verification.status,
                "verified_by": current_user.id,
                "verified_at": get_malaysia_time().isoformat()
            }
        }
    )
    
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Checklist not found")
    
    return {"message": "Checklist verified successfully"}

# Course Feedback Routes
# Feedback Template Routes
@api_router.post("/feedback-templates", response_model=FeedbackTemplate)
async def create_feedback_template(template_data: FeedbackTemplateCreate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can create feedback templates")
    
    # Delete existing template for this program
    await db.feedback_templates.delete_many({"program_id": template_data.program_id})
    
    template_obj = FeedbackTemplate(
        program_id=template_data.program_id,
        questions=template_data.questions
    )
    
    doc = template_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.feedback_templates.insert_one(doc)
    
    return template_obj

@api_router.get("/feedback-templates/program/{program_id}")
async def get_feedback_template(program_id: str, current_user: User = Depends(get_current_user)):
    template = await db.feedback_templates.find_one({"program_id": program_id}, {"_id": 0})
    if not template:
        # Return default template instead of error
        return {
            "program_id": program_id,
            "questions": [
                {"question": "Overall Training Experience", "type": "rating", "required": True},
                {"question": "Training Content Quality", "type": "rating", "required": True},
                {"question": "Trainer Effectiveness", "type": "rating", "required": True},
                {"question": "Venue & Facilities", "type": "rating", "required": True},
                {"question": "Suggestions for Improvement", "type": "text", "required": False},
                {"question": "Additional Comments", "type": "text", "required": False}
            ]
        }
    
    if isinstance(template.get('created_at'), str):
        template['created_at'] = datetime.fromisoformat(template['created_at'])
    
    return template

@api_router.delete("/feedback-templates/{template_id}")
async def delete_feedback_template(template_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can delete feedback templates")
    
    result = await db.feedback_templates.delete_one({"id": template_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Feedback template not found")
    
    return {"message": "Feedback template deleted successfully"}

@api_router.post("/feedback/submit", response_model=CourseFeedback)
async def submit_feedback(feedback_data: FeedbackSubmit, current_user: User = Depends(get_current_user)):
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can submit feedback")
    
    feedback_obj = CourseFeedback(
        participant_id=current_user.id,
        session_id=feedback_data.session_id,
        program_id=feedback_data.program_id,
        responses=feedback_data.responses
    )
    
    doc = feedback_obj.model_dump()
    doc['submitted_at'] = doc['submitted_at'].isoformat()
    
    await db.course_feedback.insert_one(doc)
    
    # Ensure participant_access record exists and update feedback status
    await db.participant_access.update_one(
        {"participant_id": current_user.id, "session_id": feedback_data.session_id},
        {"$set": {"feedback_submitted": True}},
        upsert=True
    )
    
    return feedback_obj

@api_router.get("/feedback/session/{session_id}", response_model=List[CourseFeedback])
async def get_session_feedback(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role not in ["admin", "supervisor", "coordinator"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    feedback = await db.course_feedback.find({"session_id": session_id}, {"_id": 0}).to_list(100)
    for fb in feedback:
        if isinstance(fb.get('submitted_at'), str):
            fb['submitted_at'] = datetime.fromisoformat(fb['submitted_at'])
    return feedback

@api_router.get("/feedback/company/{company_id}")
async def get_company_feedback(company_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can view company feedback")
    
    sessions = await db.sessions.find({"company_id": company_id}, {"_id": 0}).to_list(1000)
    session_ids = [s['id'] for s in sessions]
    
    feedback = await db.course_feedback.find({"session_id": {"$in": session_ids}}, {"_id": 0}).to_list(1000)
    for fb in feedback:
        if isinstance(fb.get('submitted_at'), str):
            fb['submitted_at'] = datetime.fromisoformat(fb['submitted_at'])
    
    return feedback



# Coordinator & Chief Trainer Feedback Routes

# Get Coordinator Feedback Template
@api_router.get("/coordinator-feedback-template")
async def get_coordinator_feedback_template(current_user: User = Depends(get_current_user)):
    """Get coordinator feedback template"""
    template = await db.feedback_templates.find_one({"id": "coordinator_feedback_template"}, {"_id": 0})
    if not template:
        # Create default template
        default_template = CoordinatorFeedbackTemplate()
        doc = default_template.model_dump()
        doc['updated_at'] = doc['updated_at'].isoformat()
        await db.feedback_templates.insert_one(doc)
        return default_template
    return template

# Update Coordinator Feedback Template (Admin only)
@api_router.put("/coordinator-feedback-template")
async def update_coordinator_feedback_template(
    template_update: FeedbackTemplateUpdate, 
    current_user: User = Depends(get_current_user)
):
    """Update coordinator feedback template (admin only)"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can update feedback templates")
    
    await db.feedback_templates.update_one(
        {"id": "coordinator_feedback_template"},
        {
            "$set": {
                "questions": template_update.questions,
                "updated_at": get_malaysia_time().isoformat()
            }
        },
        upsert=True
    )
    return {"message": "Template updated successfully"}

# Get Chief Trainer Feedback Template
@api_router.get("/chief-trainer-feedback-template")
async def get_chief_trainer_feedback_template(current_user: User = Depends(get_current_user)):
    """Get chief trainer feedback template"""
    template = await db.feedback_templates.find_one({"id": "chief_trainer_feedback_template"}, {"_id": 0})
    if not template:
        # Create default template
        default_template = ChiefTrainerFeedbackTemplate()
        doc = default_template.model_dump()
        doc['updated_at'] = doc['updated_at'].isoformat()
        await db.feedback_templates.insert_one(doc)
        return default_template
    return template

# Update Chief Trainer Feedback Template (Admin only)
@api_router.put("/chief-trainer-feedback-template")
async def update_chief_trainer_feedback_template(
    template_update: FeedbackTemplateUpdate, 
    current_user: User = Depends(get_current_user)
):
    """Update chief trainer feedback template (admin only)"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can update feedback templates")
    
    await db.feedback_templates.update_one(
        {"id": "chief_trainer_feedback_template"},
        {
            "$set": {
                "questions": template_update.questions,
                "updated_at": get_malaysia_time().isoformat()
            }
        },
        upsert=True
    )
    return {"message": "Template updated successfully"}

# Submit Coordinator Feedback
@api_router.post("/coordinator-feedback/{session_id}")
async def submit_coordinator_feedback(
    session_id: str,
    responses: dict,
    current_user: User = Depends(get_current_user)
):
    """Submit coordinator feedback for a session"""
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Only coordinators and admins can submit coordinator feedback")
    
    # Check if feedback already exists
    existing = await db.coordinator_feedback.find_one({"session_id": session_id}, {"_id": 0})
    
    feedback = CoordinatorFeedback(
        session_id=session_id,
        coordinator_id=current_user.id,
        responses=responses
    )
    
    doc = feedback.model_dump()
    doc['submitted_at'] = doc['submitted_at'].isoformat()
    
    if existing:
        # Update existing feedback
        await db.coordinator_feedback.update_one(
            {"session_id": session_id},
            {"$set": doc}
        )
    else:
        # Insert new feedback
        await db.coordinator_feedback.insert_one(doc)
    
    return {"message": "Coordinator feedback submitted successfully", "feedback": feedback}

# Get Coordinator Feedback for Session
@api_router.get("/coordinator-feedback/{session_id}")
async def get_coordinator_feedback(session_id: str, current_user: User = Depends(get_current_user)):
    """Get coordinator feedback for a session"""
    feedback = await db.coordinator_feedback.find_one({"session_id": session_id}, {"_id": 0})
    if not feedback:
        return None
    return feedback

# Submit Chief Trainer Feedback
@api_router.post("/chief-trainer-feedback/{session_id}")
async def submit_chief_trainer_feedback(
    session_id: str,
    responses: dict,
    current_user: User = Depends(get_current_user)
):
    """Submit chief trainer feedback for a session"""
    if current_user.role not in ["chief_trainer", "trainer", "admin"]:
        raise HTTPException(status_code=403, detail="Only trainers and admins can submit chief trainer feedback")
    
    # Check if feedback already exists
    existing = await db.chief_trainer_feedback.find_one({"session_id": session_id}, {"_id": 0})
    
    feedback = ChiefTrainerFeedback(
        session_id=session_id,
        trainer_id=current_user.id,
        responses=responses
    )
    
    doc = feedback.model_dump()
    doc['submitted_at'] = doc['submitted_at'].isoformat()
    
    if existing:
        # Update existing feedback
        await db.chief_trainer_feedback.update_one(
            {"session_id": session_id},
            {"$set": doc}
        )
    else:
        # Insert new feedback
        await db.chief_trainer_feedback.insert_one(doc)
    
    return {"message": "Chief trainer feedback submitted successfully", "feedback": feedback}

# Get Chief Trainer Feedback for Session
@api_router.get("/chief-trainer-feedback/{session_id}")
async def get_chief_trainer_feedback(session_id: str, current_user: User = Depends(get_current_user)):
    """Get chief trainer feedback for a session"""
    feedback = await db.chief_trainer_feedback.find_one({"session_id": session_id}, {"_id": 0})
    if not feedback:
        return None
    return feedback

# Certificate Routes
@api_router.get("/certificates/participant/{participant_id}", response_model=List[Certificate])
async def get_participant_certificates(participant_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role == "participant" and current_user.id != participant_id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    certificates = await db.certificates.find({"participant_id": participant_id}, {"_id": 0}).to_list(100)
    for cert in certificates:
        if isinstance(cert.get('issue_date'), str):
            cert['issue_date'] = datetime.fromisoformat(cert['issue_date'])
    return certificates

# Settings Routes
@api_router.get("/settings", response_model=Settings)
async def get_settings():
    settings = await db.settings.find_one({"id": "app_settings"}, {"_id": 0})
    if not settings:
        default_settings = Settings()
        doc = default_settings.model_dump()
        doc['updated_at'] = doc['updated_at'].isoformat()
        await db.settings.insert_one(doc)
        return default_settings
    
    if isinstance(settings.get('updated_at'), str):
        settings['updated_at'] = datetime.fromisoformat(settings['updated_at'])
    return Settings(**settings)

@api_router.post("/settings/upload-logo")
async def upload_logo(file: UploadFile = File(...), current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can update settings")
    
    file_ext = file.filename.split(".")[-1]
    filename = f"logo.{file_ext}"
    file_path = LOGO_DIR / filename
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    logo_url = f"/api/static/logos/{filename}"
    
    await db.settings.update_one(
        {"id": "app_settings"},
        {"$set": {"logo_url": logo_url, "updated_at": get_malaysia_time().isoformat()}},
        upsert=True
    )
    
    return {"logo_url": logo_url}

@api_router.put("/settings", response_model=Settings)
async def update_settings(settings_data: SettingsUpdate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can update settings")
    
    update_data = {k: v for k, v in settings_data.model_dump().items() if v is not None}
    update_data['updated_at'] = get_malaysia_time().isoformat()
    
    await db.settings.update_one(
        {"id": "app_settings"},
        {"$set": update_data},
        upsert=True
    )
    
    settings = await db.settings.find_one({"id": "app_settings"}, {"_id": 0})
    if isinstance(settings.get('updated_at'), str):
        settings['updated_at'] = datetime.fromisoformat(settings['updated_at'])
    return Settings(**settings)

# Certificate Template Upload
@api_router.post("/settings/upload-certificate-template")
async def upload_certificate_template(file: UploadFile = File(...), current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can upload templates")
    
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")
    
    filename = "certificate_template.docx"
    file_path = TEMPLATE_DIR / filename
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    template_url = f"/api/static/templates/{filename}"
    
    await db.settings.update_one(
        {"id": "app_settings"},
        {"$set": {"certificate_template_url": template_url, "updated_at": get_malaysia_time().isoformat()}},
        upsert=True
    )
    
    return {"template_url": template_url, "message": "Certificate template uploaded successfully"}

# Upload Certificate for Participant
@api_router.post("/certificates/upload/{session_id}/{participant_id}")
async def upload_participant_certificate(
    session_id: str, 
    participant_id: str,
    file: UploadFile = File(...), 
    current_user: User = Depends(get_current_user)
):
    """Upload certificate PDF for a specific participant in a session."""
    # Only coordinators assigned to the session or admins can upload
    if current_user.role == "coordinator":
        session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")
        if session.get("coordinator_id") != current_user.id:
            raise HTTPException(status_code=403, detail="You can only upload certificates for your assigned sessions")
    elif current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only coordinators and admins can upload certificates")
    
    # Validate file type
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are accepted")
    
    # Get max file size from settings
    settings = await db.settings.find_one({"id": "app_settings"}, {"_id": 0})
    max_size_mb = settings.get('max_certificate_file_size_mb', 5) if settings else 5
    max_size_bytes = max_size_mb * 1024 * 1024
    
    # Check file size
    file.file.seek(0, 2)  # Seek to end
    file_size = file.file.tell()
    file.file.seek(0)  # Reset to beginning
    
    if file_size > max_size_bytes:
        raise HTTPException(
            status_code=400, 
            detail=f"File size exceeds maximum allowed size of {max_size_mb}MB"
        )
    
    # Create unique filename
    file_extension = ".pdf"
    unique_filename = f"{session_id}_{participant_id}_{uuid.uuid4().hex[:8]}{file_extension}"
    file_path = CERTIFICATE_PDF_DIR / unique_filename
    
    # Save file
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    certificate_url = f"/api/static/certificates_pdf/{unique_filename}"
    
    # Update participant access record with certificate info
    await db.participant_access.update_one(
        {"participant_id": participant_id, "session_id": session_id},
        {
            "$set": {
                "certificate_url": certificate_url,
                "certificate_uploaded_at": get_malaysia_time().isoformat(),
                "certificate_uploaded_by": current_user.id
            }
        },
        upsert=True
    )
    
    return {
        "certificate_url": certificate_url,
        "message": "Certificate uploaded successfully",
        "file_size_mb": round(file_size / (1024 * 1024), 2)
    }

# Download Certificate for Participant
@api_router.get("/certificates/download/{session_id}/{participant_id}")
async def download_participant_certificate(
    session_id: str, 
    participant_id: str, 
    current_user: User = Depends(get_current_user)
):
    """Download certificate for a participant. Only accessible if participant has submitted feedback and clocked out."""
    
    # Check if user is the participant or admin/coordinator
    if current_user.id != participant_id and current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    # Get session
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Check if session is active (participants can only access if session is active)
    if current_user.id == participant_id and session.get("status") != "active":
        raise HTTPException(status_code=403, detail="Certificate access is not available. Session is not active.")
    
    # Get participant access
    access = await db.participant_access.find_one(
        {"participant_id": participant_id, "session_id": session_id},
        {"_id": 0}
    )
    
    if not access:
        raise HTTPException(status_code=404, detail="No certificate found")
    
    # Check if certificate exists
    certificate_url = access.get('certificate_url')
    if not certificate_url:
        raise HTTPException(status_code=404, detail="No certificate uploaded for this participant")
    
    # For participants, check eligibility (feedback + clock out)
    if current_user.id == participant_id:
        # Check feedback submission
        if not access.get('feedback_submitted', False):
            raise HTTPException(
                status_code=403, 
                detail="Certificate not available. Please submit your feedback first."
            )
        
        # Check if clocked out
        attendance = await db.attendance.find_one(
            {
                "participant_id": participant_id,
                "session_id": session_id,
                "clock_out": {"$ne": None}
            },
            {"_id": 0}
        )
        
        if not attendance:
            raise HTTPException(
                status_code=403,
                detail="Certificate not available. Please clock out first."
            )
    
    # Get file path
    filename = certificate_url.split('/')[-1]
    file_path = CERTIFICATE_PDF_DIR / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Certificate file not found")
    
    # Get participant name for filename
    participant = await db.users.find_one({"id": participant_id}, {"_id": 0})
    participant_name = participant.get('full_name', 'participant').replace(' ', '_') if participant else 'participant'
    
    return FileResponse(
        file_path,
        media_type="application/pdf",
        filename=f"{participant_name}_certificate.pdf"
    )

# Check Certificate Eligibility
@api_router.get("/certificates/eligibility/{session_id}/{participant_id}")
async def check_certificate_eligibility(
    session_id: str,
    participant_id: str,
    current_user: User = Depends(get_current_user)
):
    """Check if participant is eligible to view certificate."""
    
    # Only the participant themselves or admin/coordinator can check
    if current_user.id != participant_id and current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    # Get session
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Get participant access
    access = await db.participant_access.find_one(
        {"participant_id": participant_id, "session_id": session_id},
        {"_id": 0}
    )
    
    # Check conditions
    has_certificate = bool(access and access.get('certificate_url'))
    feedback_submitted = bool(access and access.get('feedback_submitted', False))
    
    # Check clock out
    attendance = await db.attendance.find_one(
        {
            "participant_id": participant_id,
            "session_id": session_id,
            "clock_out": {"$ne": None}
        },
        {"_id": 0}
    )
    clocked_out = bool(attendance)
    
    session_active = session.get("status") == "active"
    
    eligible = has_certificate and feedback_submitted and clocked_out and session_active
    
    return {
        "eligible": eligible,
        "has_certificate": has_certificate,
        "feedback_submitted": feedback_submitted,
        "clocked_out": clocked_out,
        "session_active": session_active,
        "certificate_url": access.get('certificate_url') if access else None,
        "message": "Eligible to download certificate" if eligible else "Not yet eligible for certificate"
    }


# Get All Certificates (Admin Only)
@api_router.get("/certificates/repository")
async def get_certificates_repository(current_user: User = Depends(get_current_user)):
    """Get all uploaded certificates for admin repository."""
    
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can access certificate repository")
    
    # Get all participant access records that have certificates
    certificates = await db.participant_access.find(
        {"certificate_url": {"$exists": True, "$ne": None}},
        {"_id": 0}
    ).to_list(length=None)
    
    # Enrich with participant, session, and program details
    enriched_certificates = []
    
    for cert in certificates:
        participant_id = cert.get('participant_id')
        session_id = cert.get('session_id')
        
        # Get participant details
        participant = await db.users.find_one({"id": participant_id}, {"_id": 0})
        
        # Get session details
        session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
        
        # Get program details if session has program_id
        program = None
        if session and session.get('program_id'):
            program = await db.programs.find_one({"id": session['program_id']}, {"_id": 0})
        
        # Get company details if session has company_id
        company = None
        if session and session.get('company_id'):
            company = await db.companies.find_one({"id": session['company_id']}, {"_id": 0})
        
        enriched_certificates.append({
            "certificate_url": cert.get('certificate_url'),
            "uploaded_at": cert.get('certificate_uploaded_at'),
            "uploaded_by": cert.get('certificate_uploaded_by'),
            "participant_id": participant_id,
            "participant_name": participant.get('full_name') if participant else 'Unknown',
            "participant_id_number": participant.get('id_number') if participant else 'N/A',
            "participant_email": participant.get('email') if participant else 'N/A',
            "session_id": session_id,
            "session_name": session.get('name') if session else 'Unknown Session',
            "session_start_date": session.get('start_date') if session else None,
            "session_end_date": session.get('end_date') if session else None,
            "program_name": program.get('name') if program else 'N/A',
            "company_name": company.get('name') if company else 'N/A',
            "feedback_submitted": cert.get('feedback_submitted', False),
        })
    
    # Sort by upload date (most recent first)
    enriched_certificates.sort(key=lambda x: x.get('uploaded_at') or '', reverse=True)
    
    return enriched_certificates


# Generate Certificate
@api_router.post("/certificates/generate/{session_id}/{participant_id}")
async def generate_certificate(session_id: str, participant_id: str, current_user: User = Depends(get_current_user)):
    # Only admin can generate, or participant can generate their own
    if current_user.role != "admin" and current_user.id != participant_id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    # Check if feedback is submitted (required for certificate)
    access = await db.participant_access.find_one(
        {"participant_id": participant_id, "session_id": session_id},
        {"_id": 0}
    )
    
    if not access:
        # Auto-create if doesn't exist
        access = await get_or_create_participant_access(participant_id, session_id)
    
    if not access.get('feedback_submitted', False):
        raise HTTPException(status_code=400, detail="Please submit feedback first. Go to your dashboard and click 'Submit Feedback' button.")
    
    # Get participant details
    participant = await db.users.find_one({"id": participant_id}, {"_id": 0})
    if not participant:
        raise HTTPException(status_code=404, detail="Participant not found")
    
    # Get session details
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Get program details
    program = await db.programs.find_one({"id": session['program_id']}, {"_id": 0})
    program_name = program['name'] if program else "Training Program"
    
    # Get company details
    company = await db.companies.find_one({"id": session['company_id']}, {"_id": 0})
    company_name = company['name'] if company else ""
    
    # Get settings for company name (already in template, no replacement needed)
    
    # Load template
    template_path = TEMPLATE_DIR / "certificate_template.docx"
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="Certificate template not found. Please upload a template first.")
    
    # Create document from template
    doc = Document(template_path)
    
    # Replace placeholders in paragraphs
    replacements = {
        '«PARTICIPANT_NAME»': participant['full_name'],
        '«IC_NUMBER»': participant['id_number'],
        '«COMPANY_NAME»': company_name,
        '«PROGRAMME NAME»': program_name,
        '<<PROGRAMME NAME>>': program_name,
        '«VENUE»': session['location'],
        '«DATE»': session['end_date']
    }
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    
    # Save as new DOCX document
    cert_filename = f"certificate_{participant_id}_{session_id}.docx"
    cert_path = CERTIFICATE_DIR / cert_filename
    doc.save(cert_path)
    
    # Convert to PDF
    pdf_filename = f"certificate_{participant_id}_{session_id}.pdf"
    pdf_path = CERTIFICATE_PDF_DIR / pdf_filename
    
    # Convert and verify
    conversion_success = convert_docx_to_pdf(cert_path, pdf_path)
    if not conversion_success or not pdf_path.exists():
        raise HTTPException(status_code=500, detail="Failed to convert certificate to PDF. Please contact support.")
    
    # Store certificate record (using PDF URL)
    cert_url = f"/api/static/certificates_pdf/{pdf_filename}"
    
    # Check if certificate already exists
    existing_cert = await db.certificates.find_one({
        "participant_id": participant_id,
        "session_id": session_id
    }, {"_id": 0})
    
    if existing_cert:
        # Update existing
        await db.certificates.update_one(
            {"id": existing_cert['id']},
            {"$set": {
                "certificate_url": cert_url,
                "issue_date": get_malaysia_time().isoformat()
            }}
        )
        cert_id = existing_cert['id']
    else:
        # Create new
        cert_obj = Certificate(
            participant_id=participant_id,
            session_id=session_id,
            program_name=program_name,
            certificate_url=cert_url
        )
        doc_cert = cert_obj.model_dump()
        doc_cert['issue_date'] = doc_cert['issue_date'].isoformat()
        await db.certificates.insert_one(doc_cert)
        cert_id = cert_obj.id
    
    return {
        "certificate_id": cert_id,
        "certificate_url": cert_url,
        "download_url": f"/api/certificates/download/{cert_id}",
        "message": "Certificate generated successfully"
    }

@api_router.get("/certificates/download/{certificate_id}")
async def download_certificate(certificate_id: str, current_user: User = Depends(get_current_user)):
    cert = await db.certificates.find_one({"id": certificate_id}, {"_id": 0})
    if not cert:
        raise HTTPException(status_code=404, detail="Certificate not found")
    
    # Only participant or admin can download
    if current_user.role != "admin" and current_user.id != cert['participant_id']:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    cert_url = cert['certificate_url']
    filename = cert_url.split('/')[-1]
    
    # Check if it's a PDF or DOCX
    if filename.endswith('.pdf'):
        file_path = CERTIFICATE_PDF_DIR / filename
        media_type = 'application/pdf'
    else:
        file_path = CERTIFICATE_DIR / filename
        media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Certificate file not found")
    
    return FileResponse(
        file_path,
        media_type=media_type,
        filename=filename,
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@api_router.get("/certificates/preview/{certificate_id}")
async def preview_certificate(certificate_id: str, current_user: User = Depends(get_current_user)):
    cert = await db.certificates.find_one({"id": certificate_id}, {"_id": 0})
    if not cert:
        raise HTTPException(status_code=404, detail="Certificate not found")
    
    # Only participant or admin can preview
    if current_user.role != "admin" and current_user.id != cert['participant_id']:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    cert_url = cert['certificate_url']
    filename = cert_url.split('/')[-1]
    
    # Check if it's a PDF or DOCX
    if filename.endswith('.pdf'):
        file_path = CERTIFICATE_PDF_DIR / filename
        media_type = 'application/pdf'
    else:
        file_path = CERTIFICATE_DIR / filename
        media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Certificate file not found")
    
    # Return PDF with inline disposition for browser preview
    return FileResponse(
        file_path,
        media_type=media_type,
        headers={"Content-Disposition": f"inline; filename={filename}"}
    )

# Static files
@api_router.get("/static/logos/{filename}")
async def get_logo(filename: str):
    file_path = LOGO_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Logo not found")
    return FileResponse(file_path)

@api_router.get("/static/certificates/{filename}")
async def get_certificate(filename: str):
    file_path = CERTIFICATE_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Certificate not found")
    return FileResponse(file_path)

@api_router.get("/static/certificates_pdf/{filename}")
async def get_certificate_pdf(filename: str):
    file_path = CERTIFICATE_PDF_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Certificate PDF not found")
    return FileResponse(
        file_path, 
        media_type='application/pdf',
        headers={
            "Content-Disposition": f"attachment; filename={filename}",
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "GET",
            "X-Content-Type-Options": "nosniff"
        }
    )

@api_router.get("/static/templates/{filename}")
async def get_template(filename: str):
    file_path = TEMPLATE_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")
    return FileResponse(file_path)

@api_router.post("/checklist-photos/upload")
async def upload_checklist_photo(file: UploadFile = File(...), current_user: User = Depends(get_current_user)):
    if current_user.role != "trainer":
        raise HTTPException(status_code=403, detail="Only trainers can upload checklist photos")
    
    if not file.content_type.startswith('image/'):
        raise HTTPException(status_code=400, detail="Only image files are allowed")
    
    # Generate unique filename
    file_extension = file.filename.split('.')[-1]
    filename = f"{str(uuid.uuid4())}.{file_extension}"
    file_path = CHECKLIST_PHOTOS_DIR / filename
    
    # Save file
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    photo_url = f"/api/static/checklist-photos/{filename}"
    return {"photo_url": photo_url}

@api_router.get("/static/checklist-photos/{filename}")
async def get_checklist_photo(filename: str):
    file_path = CHECKLIST_PHOTOS_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Photo not found")
    return FileResponse(file_path)

# ============ AI REPORT GENERATION ============

async def generate_training_report_content(session_id: str, program_id: str, company_id: str) -> str:
    """Generate comprehensive training report using GPT-5"""
    
    # Gather all data
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    program = await db.programs.find_one({"id": program_id}, {"_id": 0})
    company = await db.companies.find_one({"id": company_id}, {"_id": 0})
    
    # Get all participants
    participant_ids = session.get('participant_ids', [])
    participants = []
    for pid in participant_ids:
        user = await db.users.find_one({"id": pid}, {"_id": 0})
        if user:
            participants.append(user)
    
    # Get pre-test results
    pre_tests = await db.test_results.find({
        "session_id": session_id,
        "test_type": "pre"
    }, {"_id": 0}).to_list(100)
    
    # Get post-test results
    post_tests = await db.test_results.find({
        "session_id": session_id,
        "test_type": "post"
    }, {"_id": 0}).to_list(100)
    
    # Get checklists
    checklists = await db.vehicle_checklists.find({
        "session_id": session_id
    }, {"_id": 0}).to_list(100)
    
    # Get feedback
    feedbacks = await db.course_feedback.find({
        "session_id": session_id
    }, {"_id": 0}).to_list(100)
    
    # Get attendance
    attendance = await db.attendance_records.find({
        "session_id": session_id
    }, {"_id": 0}).to_list(100)
    
    # Create participant ID to name mapping
    participant_map = {p.get('id'): p.get('full_name') for p in participants}
    
    # Build comprehensive data structure
    training_data = {
        "session": {
            "name": session.get('name'),
            "location": session.get('location'),
            "start_date": str(session.get('start_date')),
            "end_date": str(session.get('end_date'))
        },
        "program": {
            "name": program.get('name'),
            "description": program.get('description', '')
        },
        "company": {
            "name": company.get('name')
        },
        "participants": {
            "total": len(participants),
            "names": [p.get('full_name') for p in participants],
            "id_map": participant_map
        },
        "pre_test_results": {
            "total_participants": len(pre_tests),
            "average_score": sum([t.get('score', 0) for t in pre_tests]) / len(pre_tests) if pre_tests else 0,
            "pass_rate": sum([1 for t in pre_tests if t.get('passed', False)]) / len(pre_tests) * 100 if pre_tests else 0,
            "details": [{"participant": t.get('participant_id'), "score": t.get('score'), "passed": t.get('passed')} for t in pre_tests]
        },
        "post_test_results": {
            "total_participants": len(post_tests),
            "average_score": sum([t.get('score', 0) for t in post_tests]) / len(post_tests) if post_tests else 0,
            "pass_rate": sum([1 for t in post_tests if t.get('passed', False)]) / len(post_tests) * 100 if post_tests else 0,
            "improvement": (sum([t.get('score', 0) for t in post_tests]) / len(post_tests) if post_tests else 0) - (sum([t.get('score', 0) for t in pre_tests]) / len(pre_tests) if pre_tests else 0),
            "details": [{"participant": t.get('participant_id'), "score": t.get('score'), "passed": t.get('passed')} for t in post_tests]
        },
        "checklist_summary": {
            "total_checklists": len(checklists),
            "items_needing_repair": sum([len([item for item in c.get('checklist_items', []) if item.get('status') == 'needs_repair']) for c in checklists]),
            "common_issues": [],
            "details": [{"participant": c.get('participant_id'), "items": c.get('checklist_items', [])} for c in checklists]
        },
        "feedback_summary": {
            "total_responses": len(feedbacks),
            "average_ratings": {},
            "comments": [f.get('responses', {}) for f in feedbacks]
        },
        "attendance": {
            "total_records": len(attendance),
            "attendance_rate": len([a for a in attendance if a.get('clock_out_time')]) / len(attendance) * 100 if attendance else 100
        }
    }
    
    # Create prompt for GPT-5
    prompt = f"""Generate a comprehensive Defensive Driving/Riding Training Report based on the following data:

TRAINING DETAILS:
- Program: {training_data['program']['name']}
- Company: {training_data['company']['name']}
- Session: {training_data['session']['name']}
- Location: {training_data['session']['location']}
- Dates: {training_data['session']['start_date']} to {training_data['session']['end_date']}
- Total Participants: {training_data['participants']['total']}

PRE-TEST RESULTS:
- Participants Tested: {training_data['pre_test_results']['total_participants']}
- Average Score: {training_data['pre_test_results']['average_score']:.1f}%
- Pass Rate: {training_data['pre_test_results']['pass_rate']:.1f}%

POST-TEST RESULTS:
- Participants Tested: {training_data['post_test_results']['total_participants']}
- Average Score: {training_data['post_test_results']['average_score']:.1f}%
- Pass Rate: {training_data['post_test_results']['pass_rate']:.1f}%
- Improvement: {training_data['post_test_results']['improvement']:.1f}%

VEHICLE CHECKLIST FINDINGS:
- Total Checklists Completed: {training_data['checklist_summary']['total_checklists']}
- Items Needing Repair: {training_data['checklist_summary']['items_needing_repair']}

DETAILED CHECKLIST ISSUES (items marked as 'needs_repair'):
{chr(10).join([
    f"- {training_data['participants']['id_map'].get(detail['participant'], 'Unknown participant')}: " + 
    ", ".join([
        f"Item: '{item.get('item', 'Unknown item')}' | Issue: '{item.get('comments', 'No comment')}'" 
        for item in detail['items'] 
        if item.get('status') == 'needs_repair'
    ])
    for detail in training_data['checklist_summary']['details']
    if any(item.get('status') == 'needs_repair' for item in detail['items'])
]) if training_data['checklist_summary']['items_needing_repair'] > 0 else '- No items needing repair'}

FEEDBACK:
- Total Responses: {training_data['feedback_summary']['total_responses']}

ATTENDANCE:
- Attendance Rate: {training_data['attendance']['attendance_rate']:.1f}%

Generate a professional training report with the following sections:
1. Executive Summary (2-3 paragraphs)
2. Training Overview (objectives, dates, location, participants)
3. Pre-Training Assessment (detailed analysis of pre-test results)
4. Post-Training Assessment (detailed analysis of post-test results, comparison with pre-test)
5. Vehicle Inspection Findings - READ CAREFULLY: Format each issue EXACTLY as "   - **[ITEM_CATEGORY]** - [TRAINER_COMMENT]" where:
   * ITEM_CATEGORY = The vehicle part name ONLY (Helmet, Side mirror, Safety vest, Brake, Tire, Lights, etc.)
   * TRAINER_COMMENT = The full comment from the trainer describing the issue
   * You MUST extract the item category intelligently from the 'Item' field even if it contains the full description
   * Examples:
     - If Item='No sirim helmet', extract 'Helmet' as ITEM_CATEGORY
     - If Item='No side mirror', extract 'Side mirror' as ITEM_CATEGORY  
     - If Item='Worn out', you must infer from context (likely Brake or Tire)
   * Format: "   - **Helmet** - No sirim helmet"
   * Format: "   - **Side mirror** - No side mirror"
6. Participant Feedback (summary of feedback responses)
7. Key Observations and Recommendations
8. Conclusion

Use professional language, include data-driven insights, and provide actionable recommendations for the company.
Format using Markdown with proper headings and bullet points.

ABSOLUTE CRITICAL RULES FOR VEHICLE INSPECTION SECTION:
1. Each issue line MUST start with "   - **[ITEM_CATEGORY]** - [DESCRIPTION]"
2. ITEM_CATEGORY must be a clean vehicle part name extracted from the Item field:
   - "No sirim helmet" → Extract "Helmet"
   - "No side mirror" → Extract "Side mirror"
   - "No safety vest" → Extract "Safety vest"
   - "Missing" or "Need to change" → Infer from context what part it refers to
3. Use the 'Issue' field as the DESCRIPTION after the dash
4. NEVER write "undefined" or leave item unnamed
5. Be intelligent in extracting the core item name from any description"""

    # Call GPT-5
    try:
        api_key = os.getenv('EMERGENT_LLM_KEY')
        llm = LlmChat(api_key=api_key)
        
        messages = [UserMessage(content=prompt)]
        response = llm.chat(messages=messages, model="gpt-4o")
        
        return response.content
    except Exception as e:
        logging.error(f"GPT-5 report generation failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")

@api_router.post("/reports/generate")
async def generate_report(request: ReportGenerateRequest, current_user: User = Depends(get_current_user)):
    """Generate AI training report (Coordinator only)"""
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Only coordinators can generate reports")
    
    # Get session details
    session = await db.sessions.find_one({"id": request.session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Generate report content
    content = await generate_training_report_content(
        request.session_id,
        session['program_id'],
        session['company_id']
    )
    
    # Save as draft
    report = TrainingReport(
        session_id=request.session_id,
        program_id=session['program_id'],
        company_id=session['company_id'],
        generated_by=current_user.id,
        content=content,
        status="draft"
    )
    
    await db.training_reports.insert_one(report.model_dump())
    
    return report

@api_router.get("/reports/session/{session_id}")
async def get_session_report(session_id: str, current_user: User = Depends(get_current_user)):
    """Get report for session"""
    if current_user.role not in ["coordinator", "admin", "pic_supervisor"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    report = await db.training_reports.find_one({"session_id": session_id}, {"_id": 0})
    
    # If pic_supervisor, only return published reports
    if current_user.role == "pic_supervisor":
        if not report or report.get('status') != "published":
            raise HTTPException(status_code=404, detail="No published report found")
        if current_user.id not in report.get('published_to_supervisors', []):
            raise HTTPException(status_code=403, detail="Report not published to you")
    
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    return report

@api_router.put("/reports/{report_id}")
async def update_report(report_id: str, request: ReportUpdateRequest, current_user: User = Depends(get_current_user)):
    """Update report content (draft only)"""
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Only coordinators can edit reports")
    
    report = await db.training_reports.find_one({"id": report_id}, {"_id": 0})
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    if report['status'] == "published":
        raise HTTPException(status_code=400, detail="Cannot edit published report")
    
    await db.training_reports.update_one(
        {"id": report_id},
        {"$set": {"content": request.content}}
    )
    
    return {"message": "Report updated successfully"}

@api_router.post("/reports/{report_id}/publish")
async def publish_report(report_id: str, current_user: User = Depends(get_current_user)):
    """Publish report to supervisors"""
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Only coordinators can publish reports")
    
    report = await db.training_reports.find_one({"id": report_id}, {"_id": 0})
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    # Get session to find supervisors
    session = await db.sessions.find_one({"id": report['session_id']}, {"_id": 0})
    supervisor_ids = session.get('supervisor_ids', [])
    
    await db.training_reports.update_one(
        {"id": report_id},
        {"$set": {
            "status": "published",
            "published_at": datetime.now(timezone.utc),
            "published_to_supervisors": supervisor_ids
        }}
    )
    
    return {"message": "Report published successfully", "published_to": supervisor_ids}

# ============ SUPERVISOR ENDPOINTS ============

@api_router.get("/supervisor/sessions")
async def get_supervisor_sessions(current_user: User = Depends(get_current_user)):
    """Get sessions for supervisor"""
    if current_user.role != "pic_supervisor":
        raise HTTPException(status_code=403, detail="Only supervisors can access this")
    
    # Find sessions where user is listed as supervisor
    sessions = await db.sessions.find({
        "supervisor_ids": current_user.id
    }, {"_id": 0}).to_list(100)
    
    return sessions

@api_router.get("/supervisor/attendance/{session_id}")
async def get_supervisor_session_attendance(session_id: str, current_user: User = Depends(get_current_user)):
    """Get attendance for session (Supervisor)"""
    if current_user.role != "pic_supervisor":
        raise HTTPException(status_code=403, detail="Only supervisors can access this")
    
    # Verify supervisor has access to this session
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session or current_user.id not in session.get('supervisor_ids', []):
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    # Get attendance records
    attendance = await db.attendance.find({
        "session_id": session_id
    }, {"_id": 0}).to_list(100)
    
    # Get participant details
    for record in attendance:
        participant = await db.users.find_one({"id": record['participant_id']}, {"_id": 0, "password": 0})
        if participant:
            record['participant_name'] = participant.get('full_name', 'Unknown')
            record['participant_email'] = participant.get('email', '')
        else:
            record['participant_name'] = f"Participant {record['participant_id']}"
            record['participant_email'] = ''
    
    return attendance

# Include router
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@app.on_event("startup")
async def setup_admin_account():
    """Create or update admin account on startup"""
    try:
        # PERFORMANCE OPTIMIZATION: Create database indexes
        # Indexes dramatically speed up queries on large datasets
        logging.info("📊 Creating database indexes for performance optimization...")
        
        try:
            # Users collection indexes
            await db.users.create_index("id", unique=True)
            await db.users.create_index("email", unique=True)
            await db.users.create_index("role")
            await db.users.create_index([("company_id", 1), ("role", 1)])
            
            # Sessions collection indexes
            await db.sessions.create_index("id", unique=True)
            await db.sessions.create_index("program_id")
            await db.sessions.create_index("company_id")
            await db.sessions.create_index([("start_date", 1), ("end_date", 1)])
            
            # Test results collection indexes
            await db.test_results.create_index([("session_id", 1), ("participant_id", 1)])
            await db.test_results.create_index("test_type")
            
            # Attendance collection indexes
            await db.attendance.create_index([("session_id", 1), ("participant_id", 1)])
            await db.attendance.create_index([("session_id", 1), ("date", 1)])
            
            # Participant access collection indexes
            await db.participant_access.create_index([("session_id", 1), ("participant_id", 1)], unique=True)
            
            # Feedback collection indexes
            await db.course_feedback.create_index([("session_id", 1), ("participant_id", 1)])
            
            # Vehicle issues collection indexes
            await db.vehicle_issues.create_index([("session_id", 1), ("participant_id", 1)])
            
            logging.info("✅ Database indexes created successfully")
        except Exception as idx_error:
            logging.warning(f"⚠️  Index creation warning (may already exist): {str(idx_error)}")
        
        # Admin credentials from environment variables
        admin_email = os.environ.get('ADMIN_EMAIL', 'admin@example.com')
        admin_password = os.environ.get('ADMIN_PASSWORD', 'changeme123')
        admin_name = "System Administrator"
        admin_id_number = "ADMIN001"
        
        # Check if admin exists
        existing_admin = await db.users.find_one({"role": "admin"})
        
        # Hash password
        pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
        hashed_password = pwd_context.hash(admin_password)
        
        if existing_admin:
            # Update existing admin
            await db.users.update_one(
                {"role": "admin"},
                {"$set": {
                    "email": admin_email,
                    "password": hashed_password,
                    "full_name": admin_name,
                    "id_number": admin_id_number
                }}
            )
            logging.info(f"✅ Admin account updated: {admin_email}")
        else:
            # Create new admin
            admin_doc = {
                "id": str(uuid.uuid4()),
                "email": admin_email,
                "password": hashed_password,
                "full_name": admin_name,
                "id_number": admin_id_number,
                "phone_number": "",
                "role": "admin",
                "company_id": None,
                "created_at": get_malaysia_time().isoformat()
            }
            await db.users.insert_one(admin_doc)
            logging.info(f"✅ Admin account created: {admin_email}")
        
        logging.info(f"🔐 Admin credentials: {admin_email} / {admin_password}")
        
    except Exception as e:
        logging.error(f"❌ Failed to setup admin account: {str(e)}")


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
